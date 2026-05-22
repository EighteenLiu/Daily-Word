from __future__ import annotations

import argparse
import bisect
import hashlib
import re
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from convert_xls_to_xlsx import convert_with_excel, temporary_excel_picture_compression_setting


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
}

COL_LEVEL_1 = "\u0032\u7ea7\u70b9\u4f4d"
COL_LEVEL_2 = "\u0033\u7ea7\u70b9\u4f4d"
COL_LEVEL_3 = "\u0034\u7ea7\u70b9\u4f4d"
COL_PROBLEM = "\u5177\u4f53\u95ee\u9898"
CN_LEVEL_2 = "\u4e8c\u7ea7\u70b9\u4f4d"
CN_LEVEL_3 = "\u4e09\u7ea7\u70b9\u4f4d"
CN_LEVEL_4 = "\u56db\u7ea7\u70b9\u4f4d"
CN_SONG = "\u5b8b\u4f53"
CN_TITLE = "\u95ee\u9898\u7ed3\u6784\u5316\u63d0\u53d6"
CN_SOURCE = "\u6765\u6e90\u6587\u4ef6"
CN_SHEET = "\u5de5\u4f5c\u8868"
CN_UNFILLED = "\u672a\u586b\u5199"
CN_OUTPUT_SUFFIX = "\u7ed3\u6784\u5316\u63d0\u53d6"
CN_QR_CODE = "\u4e8c\u7ef4\u7801"


@dataclass
class ExtractedImage:
    row: int
    col: int
    media_name: str
    media_bytes: bytes
    suffix: str


@dataclass
class ProblemItem:
    level1: str
    level2: str
    level3: str
    problem: str
    source_row: int
    images: list[Path] = field(default_factory=list)


@dataclass
class Level3Group:
    title: str
    items: list[ProblemItem] = field(default_factory=list)


@dataclass
class Level2Group:
    title: str
    children: dict[str, Level3Group] = field(default_factory=dict)


@dataclass
class Level1Group:
    title: str
    children: dict[str, Level2Group] = field(default_factory=dict)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Convert .xls to .xlsx through Excel, then extract a structured Word document "
            "from the point hierarchy, problem text, and row-anchored images."
        )
    )
    parser.add_argument("input", type=Path, help="Source .xls or .xlsx file.")
    parser.add_argument("-o", "--output", type=Path, help="Output .docx path.")
    parser.add_argument("--xlsx", type=Path, help="Converted/intermediate .xlsx path.")
    parser.add_argument("--sheet", help="Worksheet name. Defaults to the first worksheet.")
    parser.add_argument("--header-row", type=int, default=2, help="Header row number, 1-based. Default: 2.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing .xlsx/.docx outputs.")
    parser.add_argument("--visible", action="store_true", help="Show Excel during .xls to .xlsx conversion.")
    parser.add_argument(
        "--no-registry-protection",
        action="store_true",
        help="Do not temporarily disable Excel default picture compression during conversion.",
    )
    parser.add_argument(
        "--keep-registry-setting",
        action="store_true",
        help="Keep Excel picture compression disabled after conversion.",
    )
    return parser.parse_args()


def normalize_header(value: object) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", "", text)


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    return text.strip()


def require_modules() -> tuple[object, object]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("Missing dependency: openpyxl. Install with: pip install openpyxl") from exc
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(f"Missing dependency: python-docx/docx ({exc}). Install with: pip install python-docx") from exc
    return openpyxl, docx


def ensure_xlsx(args: argparse.Namespace) -> Path:
    source = args.input.resolve()
    if source.suffix.lower() == ".xlsx":
        return source
    if source.suffix.lower() != ".xls":
        raise ValueError(f"Input must be .xls or .xlsx: {source}")

    destination = args.xlsx.resolve() if args.xlsx else source.with_suffix(".xlsx")
    if destination.exists() and not args.overwrite:
        return destination

    with temporary_excel_picture_compression_setting(
        enabled=not args.no_registry_protection,
        keep_setting=args.keep_registry_setting,
    ):
        failures = convert_with_excel(
            jobs=[(source, destination)],
            overwrite=True,
            visible=args.visible,
            verify_media=True,
        )
    if failures:
        raise RuntimeError("Failed to convert .xls to .xlsx. See Excel COM error above.")
    return destination


def rel_target_to_zip_path(base_dir: str, target: str) -> str:
    target_path = (Path(base_dir) / target).as_posix()
    parts: list[str] = []
    for part in target_path.split("/"):
        if part in ("", "."):
            continue
        if part == "..":
            if parts:
                parts.pop()
        else:
            parts.append(part)
    return "/".join(parts)


def read_relationships(archive: zipfile.ZipFile, rels_path: str) -> dict[str, str]:
    if rels_path not in archive.namelist():
        return {}
    root = ET.fromstring(archive.read(rels_path))
    relationships: dict[str, str] = {}
    for rel in root.findall("rel:Relationship", NS):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if rel_id and target:
            relationships[rel_id] = target
    return relationships


def workbook_sheet_map(archive: zipfile.ZipFile) -> dict[str, str]:
    workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
    workbook_rels = read_relationships(archive, "xl/_rels/workbook.xml.rels")
    sheets: dict[str, str] = {}
    for sheet in workbook_root.findall("main:sheets/main:sheet", NS):
        name = sheet.attrib.get("name")
        rel_id = sheet.attrib.get(f"{{{NS['r']}}}id")
        target = workbook_rels.get(rel_id or "")
        if name and target:
            sheets[name] = rel_target_to_zip_path("xl", target)
    return sheets


def sheet_name_for_openpyxl(workbook, requested: str | None) -> str:
    if requested:
        if requested not in workbook.sheetnames:
            raise ValueError(f"Worksheet not found: {requested}. Available: {', '.join(workbook.sheetnames)}")
        return requested
    return workbook.sheetnames[0]


def extract_images_from_xlsx(
    xlsx_path: Path,
    sheet_name: str,
    ignored_columns: set[int] | None = None,
) -> list[ExtractedImage]:
    images: list[ExtractedImage] = []
    ignored_columns = ignored_columns or set()
    with zipfile.ZipFile(xlsx_path) as archive:
        sheets = workbook_sheet_map(archive)
        sheet_path = sheets.get(sheet_name)
        if not sheet_path:
            return images

        sheet_rels_path = f"{Path(sheet_path).parent.as_posix()}/_rels/{Path(sheet_path).name}.rels"
        sheet_rels = read_relationships(archive, sheet_rels_path)
        sheet_root = ET.fromstring(archive.read(sheet_path))
        drawing = sheet_root.find("main:drawing", NS)
        if drawing is None:
            return images

        drawing_rel_id = drawing.attrib.get(f"{{{NS['r']}}}id")
        drawing_target = sheet_rels.get(drawing_rel_id or "")
        if not drawing_target:
            return images

        drawing_path = rel_target_to_zip_path(Path(sheet_path).parent.as_posix(), drawing_target)
        drawing_rels_path = f"{Path(drawing_path).parent.as_posix()}/_rels/{Path(drawing_path).name}.rels"
        drawing_rels = read_relationships(archive, drawing_rels_path)
        drawing_root = ET.fromstring(archive.read(drawing_path))

        for anchor_tag in ("xdr:twoCellAnchor", "xdr:oneCellAnchor"):
            for anchor in drawing_root.findall(anchor_tag, NS):
                from_node = anchor.find("xdr:from", NS)
                blip = anchor.find(".//a:blip", NS)
                if from_node is None or blip is None:
                    continue
                row_node = from_node.find("xdr:row", NS)
                col_node = from_node.find("xdr:col", NS)
                rel_id = blip.attrib.get(f"{{{NS['r']}}}embed")
                if row_node is None or col_node is None or not rel_id:
                    continue
                col = int(col_node.text or "0") + 1
                if col in ignored_columns:
                    continue
                media_target = drawing_rels.get(rel_id)
                if not media_target:
                    continue
                media_path = rel_target_to_zip_path(Path(drawing_path).parent.as_posix(), media_target)
                if media_path not in archive.namelist():
                    continue
                suffix = Path(media_path).suffix.lower() or ".bin"
                images.append(
                    ExtractedImage(
                        row=int(row_node.text or "0") + 1,
                        col=col,
                        media_name=media_path,
                        media_bytes=archive.read(media_path),
                        suffix=suffix,
                    )
                )

    return sorted(images, key=lambda item: (item.row, item.col, item.media_name))


def nearest_data_row(image_row: int, data_rows: list[int]) -> int | None:
    if not data_rows:
        return None
    index = bisect.bisect_right(data_rows, image_row)
    if index:
        return data_rows[index - 1]
    return data_rows[0]


def save_row_images(
    images: Iterable[ExtractedImage],
    data_rows: list[int],
    temp_dir: Path,
) -> dict[int, list[Path]]:
    row_images: dict[int, list[Path]] = {row: [] for row in data_rows}
    seen: set[str] = set()
    for image in images:
        row = nearest_data_row(image.row, data_rows)
        if row is None:
            continue
        digest = hashlib.sha1(image.media_bytes).hexdigest()
        unique_key = f"{row}:{digest}:{image.media_name}"
        if unique_key in seen:
            continue
        seen.add(unique_key)
        image_path = temp_dir / f"row_{row}_{len(row_images[row]) + 1}_{digest[:10]}{image.suffix}"
        image_path.write_bytes(image.media_bytes)
        row_images[row].append(image_path)
    return row_images


def find_required_columns(headers: list[str]) -> dict[str, int]:
    aliases = {
        "level1": (COL_LEVEL_1, CN_LEVEL_2),
        "level2": (COL_LEVEL_2, CN_LEVEL_3),
        "level3": (COL_LEVEL_3, CN_LEVEL_4),
        "problem": (COL_PROBLEM,),
    }
    found: dict[str, int] = {}
    for key, names in aliases.items():
        for index, header in enumerate(headers, start=1):
            if any(name in header for name in names):
                found[key] = index
                break
    missing = [names[0] for key, names in aliases.items() if key not in found]
    if missing:
        preview = ", ".join(header or "<blank>" for header in headers[:20])
        raise ValueError(f"Missing required columns in header row: {', '.join(missing)}. Header preview: {preview}")
    return found


def find_ignored_columns(sheet, header_row: int) -> set[int]:
    return {
        col
        for col in range(1, sheet.max_column + 1)
        if any(CN_QR_CODE in normalize_header(sheet.cell(row, col).value) for row in range(1, header_row + 1))
    }


def effective_headers(sheet, header_row: int) -> list[str]:
    headers: list[str] = []
    for col in range(1, sheet.max_column + 1):
        header = ""
        for row in range(header_row, 0, -1):
            header = normalize_header(sheet.cell(row, col).value)
            if header:
                break
        headers.append(header)
    return headers


def extract_items(
    xlsx_path: Path,
    sheet_name: str | None,
    header_row: int,
    temp_dir: Path,
) -> tuple[str, list[ProblemItem], int]:
    openpyxl, _ = require_modules()
    workbook = openpyxl.load_workbook(xlsx_path, data_only=False, read_only=False)
    selected_sheet = sheet_name_for_openpyxl(workbook, sheet_name)
    sheet = workbook[selected_sheet]

    headers = effective_headers(sheet, header_row)
    columns = find_required_columns(headers)
    ignored_columns = find_ignored_columns(sheet, header_row)

    raw_items: list[ProblemItem] = []
    current_l1 = ""
    current_l2 = ""
    current_l3 = ""
    for row in range(header_row + 1, sheet.max_row + 1):
        l1 = clean_text(sheet.cell(row, columns["level1"]).value) or current_l1
        l2 = clean_text(sheet.cell(row, columns["level2"]).value) or current_l2
        l3 = clean_text(sheet.cell(row, columns["level3"]).value) or current_l3
        problem = clean_text(sheet.cell(row, columns["problem"]).value)

        current_l1, current_l2, current_l3 = l1, l2, l3
        if not any((l1, l2, l3, problem)):
            continue
        if not problem:
            continue
        raw_items.append(
            ProblemItem(
                l1 or f"{CN_UNFILLED}{COL_LEVEL_1}",
                l2 or f"{CN_UNFILLED}{COL_LEVEL_2}",
                l3 or f"{CN_UNFILLED}{COL_LEVEL_3}",
                problem,
                row,
            )
        )

    images = extract_images_from_xlsx(xlsx_path, selected_sheet, ignored_columns=ignored_columns)
    row_images = save_row_images(images, [item.source_row for item in raw_items], temp_dir)
    for item in raw_items:
        item.images = row_images.get(item.source_row, [])

    return selected_sheet, raw_items, len(images)


def set_document_styles(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = CN_SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_SONG)
    normal.font.size = Pt(10.5)

    for style_name, font_size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = CN_SONG
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_SONG)
        style.font.size = Pt(font_size)
        style.font.bold = True

    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(CN_TITLE)
    run.bold = True
    run.font.size = Pt(18)


def add_problem_images(document, image_paths: list[Path]) -> None:
    if not image_paths:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    page_width = document.sections[0].page_width
    left_margin = document.sections[0].left_margin
    right_margin = document.sections[0].right_margin
    usable_width = page_width - left_margin - right_margin
    image_width = int(usable_width / 3 * 0.92)

    for start in range(0, len(image_paths), 3):
        group = image_paths[start:start + 3]
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for image_path in group:
            run = paragraph.add_run()
            add_raw_picture(run, image_path, image_width)
            run.add_text(" ")


def add_raw_picture(run, image_path: Path, width) -> None:
    from docx.oxml.shape import CT_Inline
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.parts.image import ImagePart
    from docx.shared import Emu
    from PIL import Image

    content_types = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }
    suffix = image_path.suffix.lower()
    content_type = content_types.get(suffix)
    if not content_type:
        raise RuntimeError(f"Unsupported image type for raw Word embedding: {image_path.name}")

    blob = image_path.read_bytes()
    try:
        with Image.open(image_path) as image:
            pixel_width, pixel_height = image.size
    except Exception as exc:
        raise RuntimeError(f"Cannot read image dimensions without recompression: {image_path.name}") from exc

    image_parts = run.part.package.image_parts
    digest = hashlib.sha1(blob).hexdigest()
    image_part = image_parts._get_by_sha1(digest)
    if image_part is None:
        partname = image_parts._next_image_partname(suffix.lstrip("."))
        image_part = ImagePart(partname, content_type, blob)
        image_parts.append(image_part)

    r_id = run.part.relate_to(image_part, RT.IMAGE)
    height = Emu(int(int(width) * pixel_height / pixel_width))
    inline = CT_Inline.new_pic_inline(run.part.next_id, r_id, image_path.name, width, height)
    run._r.add_drawing(inline)


def chinese_numeral(number: int) -> str:
    digits = "\u96f6\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d"
    units = ["", "\u5341", "\u767e", "\u5343"]
    if number <= 0:
        return str(number)
    if number < 10:
        return digits[number]
    if number < 20:
        return "\u5341" + (digits[number % 10] if number % 10 else "")

    chars: list[str] = []
    zero_pending = False
    text = str(number)
    length = len(text)
    for index, char in enumerate(text):
        digit = int(char)
        unit_index = length - index - 1
        if digit == 0:
            zero_pending = bool(chars)
            continue
        if zero_pending:
            chars.append(digits[0])
            zero_pending = False
        chars.append(digits[digit] + units[unit_index])
    return "".join(chars)


def build_hierarchy(items: list[ProblemItem]) -> dict[str, Level1Group]:
    hierarchy: dict[str, Level1Group] = {}
    for item in items:
        level1 = hierarchy.setdefault(item.level1, Level1Group(item.level1))
        level2 = level1.children.setdefault(item.level2, Level2Group(item.level2))
        level3 = level2.children.setdefault(item.level3, Level3Group(item.level3))
        level3.items.append(item)
    return hierarchy


def write_docx(items: list[ProblemItem], output_path: Path, source_name: str, sheet_name: str) -> None:
    _, docx = require_modules()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = docx.Document()
    set_document_styles(document)

    meta = document.add_paragraph(f"{CN_SOURCE}: {source_name}    {CN_SHEET}: {sheet_name}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hierarchy = build_hierarchy(items)
    for level1_index, level1 in enumerate(hierarchy.values(), start=1):
        document.add_heading(f"{chinese_numeral(level1_index)}\u3001{level1.title}", level=1)
        for level2_index, level2 in enumerate(level1.children.values(), start=1):
            document.add_heading(f"\uff08{chinese_numeral(level2_index)}\uff09{level2.title}", level=2)
            for level3_index, level3 in enumerate(level2.children.values(), start=1):
                document.add_heading(f"{level3_index}. {level3.title}", level=3)
                for problem_index, item in enumerate(level3.items, start=1):
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.space_after = 0
                    paragraph.add_run(f"\uff08{problem_index}\uff09{item.problem}")
                    add_problem_images(document, item.images)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    try:
        xlsx_path = ensure_xlsx(args)
        output_path = args.output.resolve() if args.output else xlsx_path.with_name(f"{xlsx_path.stem}_{CN_OUTPUT_SUFFIX}.docx")
        if output_path.exists() and not args.overwrite:
            raise FileExistsError(f"Output already exists. Use --overwrite: {output_path}")

        with tempfile.TemporaryDirectory(prefix="xls_extract_images_") as temp:
            sheet_name, items, image_count = extract_items(xlsx_path, args.sheet, args.header_row, Path(temp))
            if not items:
                raise RuntimeError("No problem rows were extracted. Check the header row and column names.")
            write_docx(items, output_path, args.input.name, sheet_name)

        embedded_count = sum(len(item.images) for item in items)
        print(f"[ok] xlsx: {xlsx_path}")
        print(f"[ok] docx: {output_path}")
        print(f"[ok] extracted rows: {len(items)}")
        print(f"[ok] extracted images: {image_count}, embedded images: {embedded_count}")
        return 0
    except Exception as exc:
        print(f"[fail] {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
