from __future__ import annotations

import re
import posixpath
import zipfile
from collections import OrderedDict, defaultdict
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Iterable, Any
import xml.etree.ElementTree as ET

try:
    from workspace_temp import temporary_directory
except ModuleNotFoundError:
    from scripts.workspace_temp import temporary_directory


COMMUNITY_CATEGORY = "居住小区、平房胡同"
RESTAURANT_CATEGORY = "餐饮单位"
SOCIAL_UNIT_CATEGORY = "社会单位"
SPECIAL_CHECK_CATEGORY = "专项检查"
SPECIAL_CHECK_NO_PROBLEM_TYPES = {"良好，未发现问题", "良好,未发现问题", "无问题", "未发现问题"}
OUTSIDE_BUCKET_POINT = "厨余、其他垃圾桶外摆检查"
COMMUNITY_SPECIAL_NO_PROBLEM_RESULTS_BY_INDICATOR = {
    "智能回收箱是否正常运行": {"正常运行"},
    "小区宣传引导": {"有宣传氛围"},
    "居民知晓垃圾分类情况": {"合格"},
    "小区公示牌": {"良好，未发现问题", "良好,未发现问题"},
    "小区内环境": {"良好，未发现问题", "良好,未发现问题"},
    "投放点宣传指引": {"有宣传氛围"},
    "投放点公示牌设置": {"有"},
    "投放点公示牌": {"良好，未发现问题", "良好,未发现问题"},
    "容器成组设置": {"良好，未发现问题", "良好,未发现问题"},
    "容器品类成组设置": {"成组配置"},
    "容器品类设置": {"成组配置"},
    "容器标识": {"良好，未发现问题", "良好,未发现问题"},
    "遮雨棚": {"有"},
    "投放点环境": {"良好，未发现问题", "良好,未发现问题"},
    "投放点环境（容器检查）": {"良好，未发现问题", "良好,未发现问题"},
    "容器检查": {"良好，未发现问题", "良好,未发现问题"},
    "桶站便利性措施": {"有，能正常使用", "有,能正常使用"},
    "早、晚高峰投放时段开盖": {"符合要求"},
    "灭蚊蝇、地面防滑设备": {"均有设置"},
    "保洁人员作业规范": {"良好，未发现问题", "良好,未发现问题"},
    "收集车辆": {"良好，未发现问题", "良好,未发现问题"},
    "可回收物体系": {"良好，未发现问题", "良好,未发现问题"},
    "大件垃圾投放点": {"有"},
    "大件垃圾投放点设置": {"良好，未发现问题", "良好,未发现问题"},
    "装修垃圾投放点": {"有"},
    "装修垃圾投放点设置": {"良好，未发现问题", "良好,未发现问题"},
    "居民自主投放情况": {"投放正确"},
}
COMMUNITY_SPECIAL_NEUTRAL_INDICATORS = {"检查时段"}

CANONICAL_STREET_ORDER = [
    "德胜街道",
    "什刹海街道",
    "西长安街街道",
    "大栅栏街道",
    "天桥街道",
    "新街口街道",
    "金融街街道",
    "椿树街道",
    "陶然亭街道",
    "展览路街道",
    "月坛街道",
    "广内街道",
    "牛街街道",
    "白纸坊街道",
    "广外街道",
]

JINJA_TOKENS = ("{{", "}}", "{%", "%}")
NON_ISSUE_FRAGMENTS = (
    "设施齐全",
    "除臭排风灭蝇均有设置",
    "均有设置",
    "有宣传氛围",
    "居民知晓垃圾分类",
    "没有智能可回收垃圾箱",
    "无智能可回收垃圾箱",
    "该小区是纯箱房小区",
    "该小区是纯厢房小区",
    "不是箱房小区",
    "正常运行",
    "良好，未发现问题",
    "良好,未发现问题",
    "已核实",
    "投放正确",
)
NON_ISSUE_PATTERNS = (
    r"^\s*(无|无问题|没问题|正常|合格|未发现问题|未见问题|无明显问题)[。；;，,\s]*$",
    r"本小区(?:今天)?(?:一共)?(?:检查)?容器(?:检查)?数量\d+个",
    r"本小区(?:今天)?(?:一共)?检查了?\d+个容器",
    r"小区(?:今天)?(?:一共)?(?:检查了?|有|共)?\d+个(?:垃圾桶|容器)",
    r"容器检查[一二三四五六七八九十百\d]+个",
    r"访问\d+人合格\d+人",
)
COUNT_CN = {
    "一": 1,
    "二": 2,
    "两": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


@dataclass(frozen=True)
class NormalizedIssue:
    label: str
    count: int = 1

    @property
    def text(self) -> str:
        return f"{self.label}{self.count}处"


class AttachmentManager:
    def __init__(self) -> None:
        self.next_no = 1

    def assign(self, obj: dict[str, Any], enabled: bool) -> None:
        if enabled:
            obj["attachment_no"] = self.next_no
            self.next_no += 1
        else:
            obj["attachment_no"] = None


def normalize_punctuation(value: object) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("，", ",").replace("；", ";").replace("：", ":")
    return re.sub(r"\s+", "", text).strip()


def is_non_issue(text: object) -> bool:
    normalized = normalize_punctuation(text)
    if not normalized:
        return True
    if normalized.isdigit():
        return True
    if any(re.search(pattern, normalized) for pattern in NON_ISSUE_PATTERNS):
        residue = normalized
        for pattern in NON_ISSUE_PATTERNS:
            residue = re.sub(pattern, "", residue)
        for fragment in NON_ISSUE_FRAGMENTS:
            residue = residue.replace(fragment, "")
        residue = residue.strip("。.;,、:（）() ")
        return residue in {"", "无", "无问题", "没问题", "未发现问题", "未见问题", "合格", "正常"}
    return any(fragment in normalized for fragment in NON_ISSUE_FRAGMENTS) and not _has_problem_keyword(normalized)


def normalize_issue(row: Any, unit_kind: str = "residential") -> list[NormalizedIssue]:
    if unit_kind == "residential":
        indicator2_raw = _clean_issue_source(getattr(row, "indicator2", ""))
        indicator3_raw = normalize_punctuation(getattr(row, "indicator3", "")).strip("。.;,、:（）() ")
        residential_existence_problems = {
            "大件垃圾投放点": "无大件垃圾投放点",
            "装修垃圾投放点": "无装修垃圾投放点",
        }
        if indicator2_raw in residential_existence_problems and indicator3_raw == "无":
            return [NormalizedIssue(label=residential_existence_problems[indicator2_raw], count=1)]
    raw_problem = _clean_issue_source(getattr(row, "problem", ""))
    if is_non_issue(raw_problem):
        raw_problem = ""
    indicator2 = _clean_issue_source(getattr(row, "indicator2", ""))
    indicator3 = _clean_issue_source(getattr(row, "indicator3", ""))
    text = "".join(part for part in (indicator2, indicator3, raw_problem) if part)
    if not text or is_non_issue(text):
        return []

    mappings = _unit_issue_mappings() if unit_kind in {"social", "catering"} else _residential_issue_mappings()
    issues: list[NormalizedIssue] = []
    seen: set[str] = set()
    for label, patterns in mappings:
        if label in seen:
            continue
        if any(re.search(pattern, text) for pattern in patterns):
            issues.append(NormalizedIssue(label=label, count=extract_count(raw_problem or text)))
            seen.add(label)
    return issues


def extract_count(text: object) -> int:
    normalized = normalize_punctuation(text)
    explicit = [int(match) for match in re.findall(r"(\d+)处", normalized)]
    explicit.extend(
        parsed
        for parsed in (_parse_chinese_count(value) for value in re.findall(r"([零〇一二两三四五六七八九十百]+)处", normalized))
        if parsed is not None
    )
    return max(explicit) if explicit else 1


def _parse_chinese_count(value: str) -> int | None:
    if value.isdigit():
        return int(value)
    if value in COUNT_CN:
        return COUNT_CN[value]
    if "百" in value:
        left, right = value.split("百", 1)
        hundred = COUNT_CN.get(left, 1 if not left else 0) * 100
        return hundred + (_parse_chinese_count(right) or 0)
    if value.startswith("十"):
        return 10 + (_parse_chinese_count(value[1:]) or 0)
    if "十" in value:
        left, right = value.split("十", 1)
        return COUNT_CN.get(left, 0) * 10 + (_parse_chinese_count(right) or 0)
    return None


def format_issue_list(issues: Iterable[dict[str, Any] | NormalizedIssue]) -> str:
    texts = []
    for issue in issues:
        if isinstance(issue, NormalizedIssue):
            text = issue.text
        else:
            text = str(issue.get("text") or f"{issue['label']}{issue.get('count', 1)}处")
        texts.append(_normalize_final_issue_text(text))
    if not texts:
        return "无问题。"
    return "；".join(f"（{index}）{text}" for index, text in enumerate(texts, start=1)) + "。"


def _normalize_final_issue_text(text: object) -> str:
    value = "" if text is None else str(text).strip()
    while True:
        cleaned = re.sub(r"^[（(]\d+[）)]\s*", "", value)
        if cleaned == value:
            break
        value = cleaned.strip()
    return value.rstrip("。；;，,、 ")


def build_garbage_daily_context(rows: list[Any], report_date: Any) -> dict[str, Any]:
    outside_bucket_by_street = group_outside_bucket_by_street(extract_outside_bucket_issues_for_daily(rows))
    context: dict[str, Any] = {
        "report_date": getattr(report_date, "value", report_date),
        "report_date_text": report_date.chinese,
        "street_count": len(_ordered_streets(rows)),
        "city_check": {
            "summary": "今日市级检查情况未更新。",
            "note": "",
            "has_attachment": False,
            "attachment_no": None,
            "table_title": "市级重点抽查情况",
        },
        "city_check_rows": [],
        "residential_streets": _build_category_streets(
            rows,
            COMMUNITY_CATEGORY,
            "residential",
            outside_bucket_by_street=outside_bucket_by_street,
        ),
        "social_streets": _build_category_streets(rows, SOCIAL_UNIT_CATEGORY, "social"),
        "catering_streets": _build_category_streets(rows, RESTAURANT_CATEGORY, "catering"),
        "special_check": build_special_check_context(rows, report_date),
        "enforcement": _default_enforcement(report_date),
        "waste_data": {
            "date_text": _previous_day_text(report_date),
            "summary": "生活垃圾清运量数据暂未更新。",
        },
    }
    manager = AttachmentManager()
    manager.assign(context["city_check"], context["city_check"]["has_attachment"])
    manager.assign(context["enforcement"], True)
    context["enforcement"]["attachment_title"] = (
        f"{_previous_day_short(report_date)}西城区《北京市生活垃圾管理条例》专线执法统计表"
    )
    return context


def build_special_check_context(rows: list[Any], report_date: Any) -> dict[str, Any]:
    report_value = _report_date_value(report_date)
    start_value, end_value = _special_check_date_range(rows, report_value)
    summary_rows = extract_special_check_rows(rows, report_date)
    topic = _special_check_topic(rows)
    return {
        "summary": (
            f"{_format_cn_date_range(start_value, end_value)}，"
            f"区垃圾分类指挥部针对各街道的{topic}开展专项检查，"
            "各街道的问题如下表所示，"
            "具体问题照片及台账已在“垃圾分类检查工作群”发布。"
        ),
        "rows": summary_rows,
        "total_community_count": sum(row["community_count"] for row in summary_rows),
        "total_problem_count": sum(row["problem_count"] for row in summary_rows),
    }


def extract_special_check_rows(rows: Iterable[Any], report_date: Any | None = None) -> list[dict[str, Any]]:
    street_places: dict[str, set[str]] = {street: set() for street in CANONICAL_STREET_ORDER}
    street_problem_counts: dict[str, int] = {street: 0 for street in CANONICAL_STREET_ORDER}
    for row in rows:
        if _clean_special_text(getattr(row, "category", "")) != SPECIAL_CHECK_CATEGORY:
            continue

        street_name = _clean_special_text(getattr(row, "street", ""))
        if street_name not in street_places:
            continue
        point_name = _clean_special_text(getattr(row, "place", ""))
        if point_name:
            street_places[street_name].add(point_name)

        if not _is_special_check_problem_row(row):
            continue

        street_problem_counts[street_name] += _special_check_problem_weight(row)

    return [
        {
            "street_name": street,
            "community_count": len(street_places[street]),
            "problem_count": street_problem_counts[street],
        }
        for street in CANONICAL_STREET_ORDER
    ]


def normalize_datetime_text(value: object) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d 00:00:00")

    text = _clean_time_source(value)
    if not text:
        return ""
    parsed = _parse_datetime_text(text)
    if parsed is not None:
        return parsed.strftime("%Y-%m-%d %H:%M:%S")
    return text


def _is_special_check_problem_row(row: Any) -> bool:
    indicator2 = _clean_special_text(getattr(row, "indicator2", ""))
    indicator3 = _clean_special_text(getattr(row, "indicator3", ""))
    issue_text = _clean_special_text(getattr(row, "problem", ""))
    if _special_issue_text_has_problem(issue_text):
        return True
    if indicator2 in COMMUNITY_SPECIAL_NEUTRAL_INDICATORS:
        return False
    if _is_special_no_problem_result(indicator2, indicator3):
        return False
    if indicator3 in SPECIAL_CHECK_NO_PROBLEM_TYPES:
        return False
    if _is_explicit_no_problem_issue_text(issue_text):
        return False
    return bool(indicator3 or issue_text)


def _special_check_problem_weight(row: Any) -> int:
    return 1


def _is_special_no_problem_result(indicator2: str, indicator3: str) -> bool:
    if not indicator3:
        return False
    normalized_indicator2 = _clean_special_text(indicator2)
    normalized_indicator3 = _clean_special_text(indicator3)
    no_problem_results = {
        _clean_special_text(value)
        for value in COMMUNITY_SPECIAL_NO_PROBLEM_RESULTS_BY_INDICATOR.get(normalized_indicator2, set())
    }
    return normalized_indicator3 in no_problem_results


def _is_explicit_no_problem_issue_text(text: str) -> bool:
    normalized = _clean_special_text(text)
    return normalized in {
        "无",
        "无问题",
        "良好，未发现问题",
        "良好,未发现问题",
        "未发现问题",
    } or normalized.endswith("无问题") or "未发现问题" in normalized


def _special_issue_text_has_problem(text: str) -> bool:
    normalized = _clean_special_text(text)
    if not normalized or _is_explicit_no_problem_issue_text(normalized):
        return False
    return any(
        marker in normalized
        for marker in (
            "不",
            "无",
            "未",
            "遮挡",
            "脏污",
            "破损",
            "错误",
            "不符",
            "不齐全",
            "不能",
            "无法",
            "反转",
            "打不开",
            "缺",
            "混投",
            "敞口",
            "遗撒",
            "滴漏",
            "非京牌",
            "外摆",
            "满冒",
        )
    )




def infer_problem_type(issue_text: object) -> str:
    text = _clean_special_text(issue_text)
    if not text:
        return ""
    for label, patterns in _residential_issue_mappings():
        if any(re.search(pattern, text) for pattern in patterns):
            return label
    return text


def _format_special_problem_type(row: Any, problem_type: str, issue_text: str) -> str:
    problem_type = _clean_special_text(problem_type)
    station_name = _special_station_name_from_problem(issue_text)
    if station_name:
        return f"{station_name}{problem_type}"
    indicator1 = _clean_special_text(getattr(row, "indicator1", ""))
    if indicator1 in {"交投点", "清洁站", "密闭式清洁站"}:
        point_name = _clean_special_text(getattr(row, "place", ""))
        if point_name:
            return f"{point_name}{problem_type}"
    return problem_type


def _special_station_name_from_problem(issue_text: object) -> str:
    text = _clean_special_text(issue_text)
    if not text:
        return ""
    match = re.search(r"([0-9A-Za-z一二三四五六七八九十]+号)(?:垃圾)?桶站", text)
    if match:
        return f"{match.group(1)}桶站"
    match = re.search(r"([0-9A-Za-z一二三四五六七八九十]+)(?:号)?(?:垃圾)?桶站", text)
    if match:
        return f"{match.group(1)}号桶站"
    return ""


def render_garbage_daily_report(template_path: Path, context: dict[str, Any], output_path: Path) -> Path:
    from docxtpl import DocxTemplate

    with temporary_directory(prefix="garbage_daily_docxtpl_") as temp_dir:
        sanitized = _sanitize_missing_media_relationships(template_path, Path(temp_dir) / "sanitized.docx")
        _validate_garbage_daily_template(sanitized, template_path)
        prepared = _prepare_template(sanitized, Path(temp_dir) / "template.docx", context)
        doc = DocxTemplate(str(prepared))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.render(context)
        doc.save(str(output_path))
    validate_rendered_docx(output_path)
    return output_path


def validate_rendered_docx(docx_path: Path) -> None:
    text = extract_all_text_from_docx(docx_path)
    for token in JINJA_TOKENS:
        if token in text:
            raise ValueError(f"模板占位符未清理：{token}")
    forbidden = (
        "详情请见附件2",
        "本小区检查容器数量",
        "本小区容器数量",
        "容器检查",
        "设施齐全",
        "除臭排风灭蝇均有设置",
        "无问题；",
    )
    for value in forbidden:
        if value in text:
            raise ValueError(f"日报包含非标准文本：{value}")
    first = text.find("（一）德胜街道")
    second = text.find("（二）什刹海街道")
    if first != -1 and second != -1 and first > second:
        raise ValueError("居住小区街道顺序错误：德胜街道应在什刹海街道之前")


def extract_all_text_from_docx(docx_path: Path) -> str:
    from docx import Document

    document = Document(docx_path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            parts.append(paragraph.text)
        for table in section.header.tables + section.footer.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _validate_garbage_daily_template(template_path: Path, original_path: Path) -> None:
    text = extract_all_text_from_docx(template_path)
    if "{{" in text or "{%" in text:
        return
    if "桶外摆" in text:
        raise RuntimeError(
            f"垃圾分类工作日报模板选择错误：{original_path} 看起来是桶外摆日报，"
            "请在“桶外摆日报”栏选择该文件，在“垃圾分类日报模板”栏选择垃圾分类工作日报模板。"
        )
    raise RuntimeError(f"垃圾分类工作日报模板缺少 Jinja 占位符，请重新选择正确模板：{original_path}")


def _sanitize_missing_media_relationships(template_path: Path, output_path: Path) -> Path:
    relationship_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    ET.register_namespace("", relationship_namespace)

    with zipfile.ZipFile(template_path) as archive:
        names = set(archive.namelist())
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as fixed:
            for item in archive.infolist():
                data = archive.read(item.filename)
                if item.filename.endswith(".rels"):
                    data = _drop_missing_media_relationships(
                        item.filename,
                        data,
                        names,
                        relationship_namespace,
                    )
                fixed.writestr(item, data)
    return output_path


def _drop_missing_media_relationships(
    rels_member: str,
    data: bytes,
    package_members: set[str],
    relationship_namespace: str,
) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data

    changed = False
    for relationship in list(root):
        rel_type = relationship.attrib.get("Type", "")
        target = relationship.attrib.get("Target", "")
        if "image" not in rel_type or not target or "://" in target:
            continue
        member = _relationship_target_member(rels_member, target)
        if member not in package_members:
            root.remove(relationship)
            changed = True

    if not changed:
        return data
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _relationship_target_member(rels_member: str, target: str) -> str:
    target = target.replace("\\", "/")
    if target.startswith("/"):
        return posixpath.normpath(target.lstrip("/"))
    if "/_rels/" not in rels_member:
        base_dir = posixpath.dirname(rels_member)
    else:
        base_dir, rels_name = rels_member.split("/_rels/", 1)
        source_name = rels_name.removesuffix(".rels")
        base_dir = posixpath.dirname(posixpath.join(base_dir, source_name))
    return posixpath.normpath(posixpath.join(base_dir, target))


def extract_outside_bucket_issues_for_daily(rows: Iterable[Any]) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    for row in rows:
        if not _is_outside_bucket_special_row(row):
            continue
        street_name = str(getattr(row, "street", "") or "").strip()
        source_text = str(getattr(row, "problem", "") or "").strip()
        if not street_name or not source_text or is_no_problem_outside_bucket_text(source_text):
            continue
        clean_text = _normalize_outside_bucket_sentence(source_text)
        daily_text = normalize_outside_bucket_daily_text(street_name, clean_text)
        if not daily_text:
            continue
        issues.append(
            {
                "street_name": street_name,
                "street": street_name,
                "text": clean_text,
                "clean_text": clean_text,
                "daily_text": daily_text,
            }
        )
    return issues


def _is_outside_bucket_special_row(row: Any) -> bool:
    category = str(getattr(row, "category", "") or "").strip()
    indicator1 = str(getattr(row, "indicator1", "") or "").strip()
    return category == OUTSIDE_BUCKET_POINT or indicator1 == OUTSIDE_BUCKET_POINT


def is_no_problem_outside_bucket_text(text: object) -> bool:
    normalized = normalize_punctuation(text)
    return normalized in {"", "无", "无问题", "未发现问题", "良好,未发现问题", "良好，未发现问题"}


def _normalize_outside_bucket_sentence(text: object) -> str:
    value = str(text or "").strip()
    value = re.sub(r"\s+", " ", value)
    if value and not value.endswith(("。", "！", "？", "；", ";")):
        value += "。"
    return value


def normalize_outside_bucket_daily_text(street_name: str, text: str) -> str:
    value = str(text or "").strip()
    value = re.sub(r"[。；;，,、\s]+$", "", value)
    street_name = str(street_name or "").strip()
    if street_name and value.startswith(street_name):
        value = value[len(street_name) :].lstrip("，,：:、 ")
    value = value.replace("，", "").replace(",", "").replace("、", "")
    value = re.sub(r"\s+", "", value)
    value = re.sub(r"(一个|1个|一处|1处)$", "", value)
    if not value:
        return ""
    if not value.endswith("1处"):
        value += "1处"
    return value


def group_outside_bucket_by_street(outside_bucket_issues: Iterable[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: defaultdict[str, list[dict[str, Any]]] = defaultdict(list)
    seen: set[tuple[str, str]] = set()
    for issue in outside_bucket_issues or []:
        street_name = str(issue.get("street_name") or issue.get("street") or "").strip()
        if not street_name:
            continue
        source_text = str(issue.get("clean_text") or issue.get("text") or "").strip()
        if not source_text:
            continue
        daily_text = _clean_outside_bucket_daily_text(
            str(issue.get("daily_text") or "").strip()
        ) or normalize_outside_bucket_daily_text(street_name, source_text)
        if not daily_text:
            continue
        key = (street_name, daily_text)
        if key in seen:
            continue
        seen.add(key)
        item = dict(issue)
        item["street_name"] = street_name
        item["street"] = street_name
        item["daily_text"] = daily_text
        item["text"] = daily_text
        item["clean_text"] = daily_text
        grouped[street_name].append(item)
    return dict(grouped)


def make_outside_bucket_summary(street_name: str, issues: Iterable[dict[str, Any]]) -> str:
    parts: list[str] = []
    for index, issue in enumerate(issues or [], start=1):
        daily_text = _clean_outside_bucket_daily_text(str(issue.get("daily_text") or "").strip())
        if not daily_text:
            daily_text = normalize_outside_bucket_daily_text(street_name, str(issue.get("clean_text") or issue.get("text") or ""))
        if daily_text:
            parts.append(f"（{index}）{daily_text}")
    if not parts:
        return ""
    return f"{street_name}存在桶外摆问题的是：" + "；".join(parts) + "。"


def _clean_outside_bucket_daily_text(text: str) -> str:
    return re.sub(r"[。；;，,、\s]+$", "", str(text or "").strip())


def parse_outside_bucket_from_street_report_text(report_text: str) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    in_section = False
    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if "桶外摆检查" in line:
            in_section = True
            continue
        if in_section and re.match(r"^[一二三四五六七八九十]+、", line):
            break
        if not in_section or "：" not in line:
            continue
        street_name, text = line.split("：", 1)
        street_name = street_name.strip()
        text = text.strip()
        daily_text = normalize_outside_bucket_daily_text(street_name, text)
        if not street_name or not daily_text:
            continue
        issues.append(
            {
                "street_name": street_name,
                "street": street_name,
                "text": text,
                "clean_text": text,
                "daily_text": daily_text,
            }
        )
    return issues


def _build_category_streets(
    rows: list[Any],
    category: str,
    unit_kind: str,
    outside_bucket_by_street: dict[str, list[dict[str, Any]]] | None = None,
) -> list[dict[str, Any]]:
    by_street: OrderedDict[str, OrderedDict[str, list[Any]]] = OrderedDict()
    for row in rows:
        if getattr(row, "category", "") != category:
            continue
        street = getattr(row, "street", "").strip()
        place = getattr(row, "place", "").strip()
        if not street or not place:
            continue
        by_street.setdefault(street, OrderedDict()).setdefault(place, []).append(row)

    streets: list[dict[str, Any]] = []
    for street in _sort_streets(by_street):
        points = []
        for place, place_rows in by_street[street].items():
            issue_counts: OrderedDict[str, int] = OrderedDict()
            for row in place_rows:
                issue_rows = normalize_issue(row, unit_kind)
                for issue in issue_rows:
                    issue_counts[issue.label] = issue_counts.get(issue.label, 0) + issue.count
            points.append(
                {
                    "name": place,
                    "issues": [
                        {"label": label, "count": count, "text": f"{label}{count}处"}
                        for label, count in _sort_issue_counts(issue_counts, unit_kind)
                    ],
                }
            )
        streets.append(
            {
                "index_cn": _chinese_numeral(len(streets) + 1),
                "name": street,
                "points": points,
                "outside_bucket_issues": list((outside_bucket_by_street or {}).get(street, [])),
                "outside_bucket_summary": make_outside_bucket_summary(
                    street,
                    (outside_bucket_by_street or {}).get(street, []),
                ),
            }
        )
    return streets


def _residential_issue_mappings() -> list[tuple[str, tuple[str, ...]]]:
    return [
        ("桶站满冒", (r"满冒",)),
        ("桶站周边不洁", (r"周边.*不洁", r"站外不洁", r"环境不洁")),
        ("高峰时段桶站未开盖", (r"(高峰|晚高峰).*未开盖", r"未开盖.*(高峰|晚高峰)")),
        ("居民自主投放不准确", (r"混投", r"投放不规范", r"投放不准确", r"投放错误")),
        ("站外摆桶", (r"站外摆桶", r"桶外摆", r"垃圾桶外摆")),
        ("无宣传氛围", (r"无宣传氛围", r"没有看到.*宣传")),
        ("无小区公示牌", (r"无小区公示牌", r"未见公示牌")),
        ("无装修垃圾投放点", (r"装修垃圾投放点无", r"无装修垃圾投放点")),
        ("大件垃圾投放点公示牌信息错误", (r"大件垃圾投放点公示牌信息错误",)),
        ("大件垃圾投放点无大件垃圾托底上门回收信息", (r"大件垃圾投放点无大件垃圾托底上门回收信息", r"无大件垃圾托底上门回收信息")),
        ("无大件垃圾投放点", (r"^大件垃圾投放点无$", r"无大件垃圾投放点")),
        ("大件垃圾投放点未设置公示牌", (r"大件垃圾投放点未设置公示牌",)),
        ("大件垃圾投放点未设置围挡或专门隔离区", (r"大件垃圾投放点未设置围挡或专门隔离区",)),
        ("大件垃圾投放点地面未作硬化处理", (r"大件垃圾投放点地面未作硬化处理",)),
        ("大件垃圾投放点附近道路运输车通行不便", (r"大件垃圾投放点附近道路运输车通行不便",)),
        ("大件垃圾投放点大件垃圾未有序码放", (r"大件垃圾投放点大件垃圾未有序码放",)),
        ("大件垃圾投放点周围存在环境卫生死角", (r"大件垃圾投放点周围存在环境卫生死角",)),
        ("大件垃圾投放点小于6平米", (r"大件垃圾投放点小于6平米",)),
        ("大件垃圾投放点内容不规范（灭火器材不合格）", (r"大件垃圾投放点内容不规范.*灭火器材不合格", r"大件垃圾投放点内容不规范.*无灭火器")),
        ("大件垃圾投放点大件与装修或生活垃圾混放", (r"大件垃圾投放点大件与装修或生活垃圾混放",)),
        ("装修垃圾投放点公示牌信息错误", (r"装修垃圾投放点公示牌信息错误",)),
        ("装修垃圾投放点未设置公示牌", (r"装修垃圾投放点未设置公示牌",)),
        ("装修垃圾投放点未设置围挡或专门隔离区", (r"装修垃圾投放点未设置围挡或专门隔离区",)),
        ("装修垃圾投放点地面未作硬化处理", (r"装修垃圾投放点地面未作硬化处理",)),
        ("装修垃圾投放点附近道路运输车通行不便", (r"装修垃圾投放点附近道路运输车通行不便",)),
        ("装修垃圾投放点未袋装并有序码放", (r"装修垃圾投放点未袋装并有序码放",)),
        ("装修垃圾投放点周边存在环境卫生死角", (r"装修垃圾投放点周边存在环境卫生死角",)),
        ("装修垃圾投放点小于6平米", (r"装修垃圾投放点小于6平米",)),
        ("装修垃圾投放点内容不规范（灭火器材不合格）", (r"装修垃圾投放点内容不规范.*灭火器材不合格", r"装修垃圾投放点内容不规范.*无灭火器")),
        ("装修垃圾投放点装修与大件或生活垃圾混放", (r"装修垃圾投放点装修与大件或生活垃圾混放",)),
        ("无装修垃圾备案", (r"无装修垃圾备案",)),
        ("无桶站", (r"无桶站", r"未见垃圾桶站")),
        ("散桶", (r"散桶",)),
        ("垃圾车混装混运", (r"垃圾车混装混运",)),
    ]


def _unit_issue_mappings() -> list[tuple[str, tuple[str, ...]]]:
    return [
        ("分类投放不准确", (r"分类投放不准确", r"投放不准确", r"投放错误", r"混投", r"厨余.*其他", r"其他.*厨余")),
        ("桶站周边不洁", (r"周边.*不洁", r"环境不洁")),
        ("桶站满冒", (r"满冒",)),
        ("站外摆桶", (r"站外摆桶", r"桶外摆", r"垃圾桶外摆")),
        ("无宣传氛围", (r"无宣传氛围", r"没有看到.*宣传")),
    ]


def _sort_issue_counts(issue_counts: OrderedDict[str, int], unit_kind: str) -> list[tuple[str, int]]:
    mappings = _unit_issue_mappings() if unit_kind in {"social", "catering"} else _residential_issue_mappings()
    order = {label: index for index, (label, _patterns) in enumerate(mappings)}
    return sorted(issue_counts.items(), key=lambda item: (order.get(item[0], len(order)), item[0]))


def _clean_issue_source(value: object) -> str:
    text = normalize_punctuation(value)
    text = re.sub(r"^[（(]?\d+[）)]?", "", text)
    for pattern in NON_ISSUE_PATTERNS:
        text = re.sub(pattern, "", text)
    for fragment in NON_ISSUE_FRAGMENTS:
        text = text.replace(fragment, "")
    text = re.sub(r"\d+个(?:垃圾桶|容器)", "", text)
    text = re.sub(r"([0-9A-Za-z]+|[一二三四五六七八九十]+)(?:号)?桶站", "", text)
    return text.strip("。.;,、:（）() ")


def _has_problem_keyword(text: str) -> bool:
    return any(
        keyword in text
        for keyword in (
            "不洁",
            "满冒",
            "未开盖",
            "混投",
            "不规范",
            "不准确",
            "桶外摆",
            "站外摆桶",
            "无宣传",
            "无小区公示牌",
            "无桶站",
            "散桶",
            "混装混运",
            "信息错误",
            "托底上门回收信息",
            "未设置",
            "未作硬化",
            "通行不便",
            "未有序",
            "未袋装",
            "卫生死角",
            "小于6平米",
            "灭火器材不合格",
            "混放",
            "无装修垃圾备案",
            "无大件垃圾投放点",
            "无装修垃圾投放点",
        )
    )


def _is_street_outside_bucket(row: Any, issues: list[NormalizedIssue]) -> bool:
    if not any(issue.label == "站外摆桶" for issue in issues):
        return False
    text = normalize_punctuation(getattr(row, "problem", ""))
    if "桶站" in text or "小区" in text:
        return False
    return any(keyword in text for keyword in ("门口", "道路", "路", "街", "胡同", "大街", "北口", "南口", "东口", "西口", "商户"))


def _format_outside_bucket_issue(row: Any) -> str:
    text = str(getattr(row, "problem", "")).strip("。；; ")
    text = re.sub(r"一处$", "1处", text)
    if not re.search(r"\d+处$", text):
        text += "1处"
    return text


def _dedupe_issue_dicts(items: list[dict[str, str]]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    seen: set[str] = set()
    for item in items:
        text = item["text"]
        if text in seen:
            continue
        seen.add(text)
        result.append(item)
    return result


def _ordered_streets(rows: list[Any]) -> list[str]:
    seen: set[str] = set()
    streets: list[str] = []
    for row in rows:
        street = getattr(row, "street", "").strip()
        if street and street not in seen:
            seen.add(street)
            streets.append(street)
    return _sort_streets({street: None for street in streets})


def _sort_streets(streets: dict[str, Any] | OrderedDict[str, Any]) -> list[str]:
    order = {street: index for index, street in enumerate(CANONICAL_STREET_ORDER)}
    return sorted(streets, key=lambda street: (order.get(street, len(order)), street))


def _default_enforcement(report_date: Any) -> dict[str, Any]:
    return {
        "attachment_no": None,
        "date_text": _previous_day_text(report_date),
        "attachment_title": "",
        "rows": [
            {"street_name": street, "check_count": 0, "case_count": 0, "fine_amount": 0}
            for street in CANONICAL_STREET_ORDER
        ],
        "total_checks": 0,
        "case_count": 0,
        "fine_amount": 0,
    }


def _prepare_template(template_path: Path, output_path: Path, context: dict[str, Any]) -> Path:
    from docx import Document

    document = Document(str(template_path))
    if not context.get("city_check", {}).get("has_attachment"):
        for table in list(document.tables):
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "city_check" in table_text:
                table_element = table._tbl
                table_element.getparent().remove(table_element)
    if not context.get("special_check", {}).get("rows"):
        for table in list(document.tables):
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "special_check.rows" in table_text:
                table_element = table._tbl
                table_element.getparent().remove(table_element)
    for paragraph in document.paragraphs:
        _replace_paragraph_text(paragraph, "附件2", "附件{{ enforcement.attachment_no }}")
    document.save(str(output_path))
    return output_path


def _replace_paragraph_text(paragraph: Any, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    replaced = paragraph.text.replace(old, new)
    for index, run in enumerate(paragraph.runs):
        run.text = replaced if index == 0 else ""


def _previous_day_text(report_date: Any) -> str:
    return _format_cn_date(_report_date_value(report_date) - timedelta(days=1))


def _previous_day_short(report_date: Any) -> str:
    value = _report_date_value(report_date) - timedelta(days=1)
    return f"{value.month}.{value.day}"


def _report_date_value(report_date: Any) -> date:
    value = getattr(report_date, "value", report_date)
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    raise TypeError(f"Unsupported report date: {report_date!r}")


def _format_cn_date(value: date) -> str:
    return f"{value.year}年{value.month}月{value.day}日"


def _format_cn_date_range(start_value: date, end_value: date) -> str:
    if start_value == end_value:
        return _format_cn_date(start_value)
    if start_value.year == end_value.year:
        if start_value.month == end_value.month:
            return f"{start_value.year}年{start_value.month}月{start_value.day}日-{end_value.day}日"
        return f"{start_value.year}年{start_value.month}月{start_value.day}日-{end_value.month}月{end_value.day}日"
    return f"{_format_cn_date(start_value)}-{_format_cn_date(end_value)}"


def _special_check_date_range(rows: Iterable[Any], report_value: date) -> tuple[date, date]:
    dates: list[date] = []
    for row in rows:
        if _clean_special_text(getattr(row, "category", "")) != SPECIAL_CHECK_CATEGORY:
            continue
        for value in (getattr(row, "created_time", ""), getattr(row, "report_time", "")):
            parsed = _parse_datetime_text(_clean_time_source(value))
            if parsed is not None:
                dates.append(parsed.date())
                break
    if dates:
        return min(dates), max(dates)
    return report_value - timedelta(days=1), report_value


def _special_check_topic(rows: Iterable[Any]) -> str:
    issue_texts: list[str] = []
    for row in rows:
        if _clean_special_text(getattr(row, "category", "")) != SPECIAL_CHECK_CATEGORY:
            continue
        if not _is_special_check_problem_row(row):
            continue
        issue_texts.append(
            "".join(
                _clean_special_text(getattr(row, field, ""))
                for field in ("indicator1", "indicator2", "indicator3", "problem")
            )
        )
    combined = "".join(issue_texts)
    if not combined:
        return "桶站建设问题"
    if any(keyword in combined for keyword in ("公示", "标识", "便利", "脚踏", "大件", "建设", "颜色", "电话")):
        return "桶站建设问题"
    if any(keyword in combined for keyword in ("满冒", "脏污", "破损", "周边不洁")):
        return "桶站满冒脏污问题"
    return "桶站建设问题"


def _clean_special_text(value: object) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[ \t\u3000]+", "", text).strip()


def _clean_time_source(value: object) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[ \t\u3000]+", " ", text).strip()


def _parse_datetime_text(text: str) -> datetime | None:
    normalized = text.strip()
    normalized = normalized.replace("年", "-").replace("月", "-").replace("日", " ")
    normalized = normalized.replace("/", "-").replace(".", "-")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    normalized = re.sub(r"^(\d{4})-(\d{1,2})-(\d{1,2})(\d{1,2}:\d{2})", r"\1-\2-\3 \4", normalized)
    formats = (
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%Y-%m-%d %H:%M:%S.%f",
    )
    for fmt in formats:
        try:
            return datetime.strptime(normalized, fmt)
        except ValueError:
            continue
    return None


def _datetime_sort_key(text: str) -> datetime:
    parsed = _parse_datetime_text(text)
    return parsed or datetime.max


def _chinese_numeral(number: int) -> str:
    digits = "零一二三四五六七八九"
    if number <= 0:
        return str(number)
    if number < 10:
        return digits[number]
    if number < 20:
        return "十" + (digits[number % 10] if number % 10 else "")
    tens, ones = divmod(number, 10)
    return digits[tens] + "十" + (digits[ones] if ones else "")
