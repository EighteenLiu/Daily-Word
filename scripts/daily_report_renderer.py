from __future__ import annotations

import re
from dataclasses import dataclass, fields, is_dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

try:
    from workspace_temp import temporary_directory
except ModuleNotFoundError:
    from scripts.workspace_temp import temporary_directory


IMAGE_WIDTH_CM = 10.2
IMAGE_HEIGHT_CM = 5.74


@dataclass(frozen=True)
class ImageCompressionProfile:
    width: int
    height: int
    quality: int


IMAGE_COMPRESSION_PROFILES: dict[str, ImageCompressionProfile | None] = {
    "none": None,
    "light": ImageCompressionProfile(width=1920, height=1080, quality=88),
    "standard": ImageCompressionProfile(width=1600, height=900, quality=82),
    "strong": ImageCompressionProfile(width=1280, height=720, quality=72),
}
DEFAULT_IMAGE_COMPRESSION = "standard"


def render_street_report_docx_from_docxtpl(
    report: Any,
    template_path: Path,
    output_path: Path,
    report_title: str = "",
    report_date_text: str = "",
    outside_bucket_issues: list[Any] | None = None,
    image_width_cm: float = IMAGE_WIDTH_CM,
    image_height_cm: float = IMAGE_HEIGHT_CM,
    image_compression: str = DEFAULT_IMAGE_COMPRESSION,
) -> Path:
    """
    使用 docxtpl 直接渲染 Word 模板。
    这个函数会保留模板的页眉、页脚、表格、样式、段落格式等。
    """
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Cm

    with temporary_directory(prefix="daily_report_docxtpl_") as temp_dir:
        prepared_template = _prepare_docxtpl_template(
            template_path,
            Path(temp_dir) / "template.docx",
        )
        doc = DocxTemplate(str(prepared_template))
        image_cache: dict[Path, Path] = {}

        def convert(obj, width_cm: float = image_width_cm, height_cm: float = image_height_cm):
            if isinstance(obj, Path):
                image_path = _word_image_path(obj, Path(temp_dir), image_cache, image_compression)
                if image_path:
                    return InlineImage(doc, str(image_path), width=Cm(width_cm), height=Cm(height_cm))
                return ""

            if isinstance(obj, list):
                return [convert(item, width_cm, height_cm) for item in obj]

            if isinstance(obj, tuple):
                return [convert(item, width_cm, height_cm) for item in obj]

            if isinstance(obj, dict):
                return {key: convert(value, width_cm, height_cm) for key, value in obj.items()}

            if is_dataclass(obj):
                converted = {
                    field.name: convert(getattr(obj, field.name), width_cm, height_cm)
                    for field in fields(obj)
                }
                if hasattr(obj, "overall_items"):
                    converted["overall_items"] = convert(getattr(obj, "overall_items"), width_cm, height_cm)
                if hasattr(obj, "image_rows"):
                    converted["image_rows"] = convert(getattr(obj, "image_rows"), width_cm, height_cm)
                if hasattr(obj, "images"):
                    converted["images"] = convert(getattr(obj, "images"), width_cm, height_cm)
                return converted

            return obj

        context = convert(report)
        context["report_title"] = report_title or "区级检查日报"
        context["report_date_text"] = report_date_text
        context["outside_bucket_issues"] = convert(outside_bucket_issues or [], 7.21, 4.06)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.render(context)
        saved_path = _save_with_available_path(lambda path: doc.save(str(path)), output_path)
    return saved_path


def _word_image_path(
    image_path: Path,
    temp_dir: Path,
    image_cache: dict[Path, Path],
    image_compression: str = DEFAULT_IMAGE_COMPRESSION,
) -> Path | None:
    if image_path in image_cache:
        return image_cache[image_path]
    if not image_path.exists():
        return None

    from PIL import Image
    from PIL import ImageOps

    profile = image_compression_profile(image_compression)
    output = temp_dir / f"image_{len(image_cache) + 1}.jpg"
    try:
        with Image.open(image_path) as image:
            image = ImageOps.exif_transpose(image)
            if image.mode in ("RGBA", "LA") or (image.mode == "P" and "transparency" in image.info):
                canvas = Image.new("RGB", image.size, (255, 255, 255))
                canvas.paste(image.convert("RGBA"), mask=image.convert("RGBA").split()[-1])
                image = canvas
            else:
                image = image.convert("RGB")
            quality = 95
            if profile is not None:
                target_size = (
                    min(profile.width, image.width),
                    min(profile.height, image.height),
                )
                if image.size != target_size:
                    image = image.resize(target_size, Image.Resampling.LANCZOS)
                quality = profile.quality
            image.save(output, "JPEG", quality=quality, optimize=True, progressive=False)
    except Exception:
        return None
    image_cache[image_path] = output
    return output


def image_compression_profile(name: str | None) -> ImageCompressionProfile | None:
    key = (name or DEFAULT_IMAGE_COMPRESSION).strip().lower()
    if key not in IMAGE_COMPRESSION_PROFILES:
        key = DEFAULT_IMAGE_COMPRESSION
    return IMAGE_COMPRESSION_PROFILES[key]


def _prepare_docxtpl_template(template_path: Path, output_path: Path) -> Path:
    document = Document(str(template_path))
    _split_soft_break_paragraphs(document)
    _ensure_community_intro_placeholder(document)
    _remove_template_instruction_tail(document)
    _drop_unmatched_jinja_control_paragraphs(document)
    document.save(str(output_path))
    return output_path


def _ensure_community_intro_placeholder(document: Document) -> None:
    intro_placeholder = "{{ community.overall_intro }}"
    problem_placeholder = "{{ community.overall_problem_summary }}"
    for paragraph in document.paragraphs:
        text = paragraph.text
        if problem_placeholder not in text or intro_placeholder in text:
            continue
        if "存在的问题是：" not in text:
            continue
        if paragraph.runs:
            paragraph.runs[0].text = intro_placeholder + paragraph.runs[0].text
        else:
            paragraph.add_run(intro_placeholder)


def _split_soft_break_paragraphs(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        lines = paragraph.text.splitlines()
        if len(lines) <= 1:
            continue

        parent = paragraph._element.getparent()
        if parent is None:
            continue

        insert_at = parent.index(paragraph._element)
        for line in lines:
            new_paragraph = _clone_paragraph_with_text(paragraph, line)
            insert_at += 1
            parent.insert(insert_at, new_paragraph._element)
        parent.remove(paragraph._element)


def _clone_paragraph_with_text(paragraph, text: str):
    from copy import deepcopy

    new_paragraph = deepcopy(paragraph)
    for run in new_paragraph.runs:
        run.text = ""
    if new_paragraph.runs:
        new_paragraph.runs[0].text = text
    else:
        new_paragraph.add_run(text)
    return new_paragraph


def _remove_template_instruction_tail(document: Document) -> None:
    marker = "模板上下文字段说明"
    paragraphs = list(document.paragraphs)
    start = next((index for index, para in enumerate(paragraphs) if marker in para.text), None)
    if start is None:
        return

    if start > 0 and _paragraph_has_page_break_only(paragraphs[start - 1]):
        start -= 1

    for paragraph in paragraphs[start:]:
        _remove_paragraph(paragraph)


def _drop_unmatched_jinja_control_paragraphs(document: Document) -> None:
    stack: list[str] = []
    for paragraph in list(document.paragraphs):
        tag = _jinja_control_tag(paragraph.text)
        if not tag:
            continue

        if tag in {"if", "for"}:
            stack.append(tag)
            continue
        if tag == "endif":
            if stack and stack[-1] == "if":
                stack.pop()
            else:
                _remove_paragraph(paragraph)
            continue
        if tag == "endfor":
            if stack and stack[-1] == "for":
                stack.pop()
            else:
                _remove_paragraph(paragraph)



def _paragraph_has_page_break_only(paragraph) -> bool:
    return not paragraph.text.strip() and bool(paragraph._p.xpath('.//*[local-name()="br"]'))


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


try:
    from scripts.daily_report_builder import StreetReport, UnitSection
except ModuleNotFoundError:
    from daily_report_builder import StreetReport, UnitSection




def render_street_report_docx_from_jinja(
    report: "StreetReport",
    template_path: Path,
    output_path: Path,
    report_title: str = "",
    report_date_text: str = "",
    max_template_paras: int | None = 73,
) -> Path:
    """Render a StreetReport using a Jinja2-based .docx template."""
    import jinja2
    
    from docx import Document as DocxDocument
    from docx.shared import Cm as DX_Cm, Pt as DX_Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as DX_ALIGN
    from docx.oxml.ns import qn as dx_qn

    # Step 1: Extract Jinja2 template text from the docx
    document = DocxDocument(str(template_path))
    template_lines = [p.text for p in document.paragraphs]
    if max_template_paras:
        template_lines = template_lines[:max_template_paras]
    template_lines = _drop_unmatched_jinja_end_tags(template_lines)
    template_text = "\n".join(template_lines)

    # Step 2: Replace all Path objects with unique marker strings
    marker_to_image: dict[str, Path] = {}
    marker_counter: list[int] = [0]

    def _replace_paths(obj: Any) -> Any:
        if isinstance(obj, Path):
            marker = f"__IMG_{marker_counter[0]}__"
            marker_to_image[marker] = obj
            marker_counter[0] += 1
            return marker
        if isinstance(obj, list):
            return [_replace_paths(item) for item in obj]
        if isinstance(obj, dict):
            return {k: _replace_paths(v) for k, v in obj.items()}
        if hasattr(obj, "__dataclass_fields__"):
            return {f: _replace_paths(getattr(obj, f)) for f in obj.__dataclass_fields__}
        return obj

    context = _replace_paths(report)
    context["report_title"] = report_title
    context["report_date_text"] = report_date_text

    # Step 3: Render with Jinja2
    env = jinja2.Environment()
    rendered = env.from_string(template_text).render(**context)

    # Step 4: Build output docx from rendered text
    clean_doc = DocxDocument()
    _set_document_styles(clean_doc)

    for line in rendered.split("\n"):
        text = line.strip()
        if not text:
            continue

        if text.startswith("__IMG_") and text.endswith("__"):
            img_path = marker_to_image.get(text)
            if img_path and img_path.exists():
                para = clean_doc.add_paragraph()
                para.alignment = DX_ALIGN.CENTER
                para.paragraph_format.space_before = DX_Pt(0)
                para.paragraph_format.space_after = DX_Pt(0)
                run = para.add_run()
                try:
                    from PIL import Image as PILImg
                    from io import BytesIO
                    with PILImg.open(img_path) as pil_img:
                        buf = BytesIO()
                        pil_img.convert('RGB').save(buf, 'PNG')
                        buf.seek(0)
                        run.add_picture(buf, width=DX_Cm(IMAGE_WIDTH_CM), height=DX_Cm(IMAGE_HEIGHT_CM))
                except Exception:
                    fb = para.add_run("[图片无法插入：" + img_path.name + "]")
                    fb.font.size = DX_Pt(9)
            else:
                fb = clean_doc.add_paragraph().add_run("[图片无法插入图像]")
            continue

        para = clean_doc.add_paragraph()
        para.paragraph_format.space_before = DX_Pt(0)
        para.paragraph_format.space_after = DX_Pt(0)
        run = para.add_run(text)
        run.font.name = FONT_FANGSONG
        run._element.rPr.rFonts.set(dx_qn("w:eastAsia"), FONT_FANGSONG)
        run.font.size = DX_Pt(14)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    return _save_with_available_path(lambda path: clean_doc.save(str(path)), output_path)


def _drop_unmatched_jinja_end_tags(lines: list[str]) -> list[str]:
    stack: list[str] = []
    filtered: list[str] = []
    for line in lines:
        tag = _jinja_control_tag(line)
        if not tag:
            filtered.append(line)
            continue

        if tag in {"if", "for"}:
            stack.append(tag)
            filtered.append(line)
            continue
        if tag == "endif":
            if stack and stack[-1] == "if":
                stack.pop()
                filtered.append(line)
            continue
        if tag == "endfor":
            if stack and stack[-1] == "for":
                stack.pop()
                filtered.append(line)
            continue
        filtered.append(line)
    return filtered


def _jinja_control_tag(text: str) -> str | None:
    stripped = text.strip()
    tag_match = re.fullmatch(r"\{%-?\s*(?:(?:p|tr|tc|r)\s+)?(\w+)\b.*?-?%\}", stripped)
    return tag_match.group(1) if tag_match else None


FONT_SONG = "\u5b8b\u4f53"
FONT_FANGSONG = "\u4eff\u5b8b"
FONT_HEITI = "\u9ed1\u4f53"


def render_street_report_docx(
    report: StreetReport,
    output_path: Path,
    outside_bucket_issues: list[Any] | None = None,
    image_compression: str = DEFAULT_IMAGE_COMPRESSION,
) -> Path:
    document = Document()
    _set_document_styles(document)

    with temporary_directory(prefix="daily_report_images_") as temp_dir:
        image_cache = ImageCache(Path(temp_dir), image_compression=image_compression)
        if report.communities:
            _add_heading(document, "一、居住小区", level=1)
            section_index = 1
            for community in report.communities:
                _add_heading(document, f"（{community.index_cn}）{community.name}", level=2)
                _add_heading(document, "1.小区整体情况", level=3)
                if community.is_pure_box_room:
                    _add_body(
                        document,
                        "该小区垃圾分类情况好的方面主要有：1.四类垃圾桶设置齐全，且容器完好、洁净；2.有开展垃圾分类宣传工作；3.站点指导员按时上岗；4.小区垃圾投放规范，无垃圾乱堆乱放、投放不规范等现象；5.厨余、可回收物和有害垃圾桶分类纯净；6.有装修垃圾和大件垃圾投放点并且规范。存在的问题是："
                        f"{community.overall_problem_summary}",
                    )
                else:
                    _add_body(document, f"{community.overall_intro}存在的问题是：{community.overall_problem_summary}")
                promo_text = community.promo_text if community.promo_text != "无问题" else ""
                notice_board_text = community.notice_board_text if community.notice_board_text != "无问题" else ""
                _add_body(document, f"（1）小区宣传氛围：{promo_text}")
                _add_images(document, community.promo_images, image_cache)
                _add_body(document, f"（2）小区公示牌：{notice_board_text}")
                _add_images(document, community.notice_board_images, image_cache)
                if community.is_pure_box_room:
                    _add_body(document, "（3）装修垃圾投放点设置")
                    _add_body(document, "预约收集，集中密闭运输。")
                    _add_body(document, "（4）大件垃圾投放点设置")
                    _add_body(document, "预约收集，集中密闭运输。")
                if community.community_litter_text:
                    _add_body(document, f"（5）小区垃圾乱堆乱放、投放不规范现象：{community.community_litter_text}")
                    _add_images(document, community.community_litter_images, image_cache)

                _add_heading(document, "2.桶站设置情况", level=3)
                if community.stations:
                    for station in community.stations:
                        _add_heading(document, f"{station.title}：{station.problem_summary}", level=3)
                        _add_images(document, station.images, image_cache)
                else:
                    _add_heading(document, "1号桶站设置情况：无问题", level=3)

                if community.resident_delivery:
                    _add_heading(document, "3.居民投放情况", level=3)
                    _add_body(document, community.resident_delivery.summary)
                    _add_images(document, community.resident_delivery.error_images, image_cache)

        section_number = 1
        if report.communities:
            section_number += 1

        if report.restaurants:
            _render_units(document, f"{_chinese_section_number(section_number)}、餐饮单位", report.restaurants, image_cache)
            section_number += 1

        if report.social_units:
            _render_units(document, f"{_chinese_section_number(section_number)}、社会单位", report.social_units, image_cache)
            section_number += 1

        if outside_bucket_issues:
            _render_outside_bucket_issues(
                document,
                f"{_chinese_section_number(section_number)}、桶外摆检查",
                outside_bucket_issues,
                image_cache,
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        saved_path = _save_with_available_path(lambda path: document.save(path), output_path)
    return saved_path


def _save_with_available_path(save_func, output_path: Path) -> Path:
    try:
        save_func(output_path)
        return output_path
    except PermissionError:
        fallback = _next_available_output_path(output_path)
        print(f"[warn] 目标文件被占用，已另存为: {fallback}")
        save_func(fallback)
        return fallback


def _next_available_output_path(output_path: Path) -> Path:
    for index in range(1, 100):
        candidate = output_path.with_name(f"{output_path.stem}（另存{index}）{output_path.suffix}")
        if not candidate.exists():
            return candidate
    raise PermissionError(f"目标文件被占用，且无法找到可用另存文件名: {output_path}")


class ImageCache:
    def __init__(self, temp_dir: Path, image_compression: str = DEFAULT_IMAGE_COMPRESSION) -> None:
        self.temp_dir = temp_dir
        self.image_compression = image_compression
        self.converted: dict[Path, Path] = {}

    def word_image_path(self, image_path: Path) -> Path:
        prepared = _word_image_path(
            image_path,
            self.temp_dir,
            self.converted,
            self.image_compression,
        )
        if prepared is None:
            raise FileNotFoundError(image_path)
        return prepared


def _render_units(document: Document, heading: str, units: list[UnitSection], image_cache: ImageCache) -> None:
    _add_heading(document, heading, level=1)
    for unit in units:
        _add_heading(document, f"（{unit.index_cn}）{unit.name}", level=2)
        _add_heading(document, "1.整体情况", level=3)
        _add_body(document, f"存在的问题是：{unit.overall_problem_summary}")
        suffix = f"{unit.promo_text}" if unit.promo_text and unit.promo_text != "无问题" else ""
        _add_body(document, f"（1）宣传氛围：{suffix}")
        _add_images(document, unit.promo_images, image_cache)
        _add_heading(document, f"2.桶站设置情况：{unit.container_problem_summary}", level=3)
        _add_images(document, unit.container_images, image_cache)


def _render_outside_bucket_issues(
    document: Document,
    heading: str,
    issues: list[Any],
    image_cache: ImageCache,
) -> None:
    _add_heading(document, heading, level=1)
    for issue in issues:
        street_name = _get_issue_value(issue, "street_name", "street")
        clean_text = _get_issue_value(issue, "clean_text", "text")
        if clean_text and not clean_text.endswith(("。", "；", ";", "！", "!", "?", "？")):
            clean_text += "。"
        _add_body(document, f"{street_name}：{clean_text}")
        image_rows = _get_issue_image_rows(issue)
        if not image_rows:
            image_rows = [_get_issue_images(issue)]
        for row in image_rows:
            if len(row) <= 1:
                _add_images(document, row, image_cache, width_cm=7.21, height_cm=4.06)
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            first = paragraph.add_run()
            try:
                first.add_picture(
                    str(image_cache.word_image_path(row[0])),
                    width=Cm(7.21),
                    height=Cm(4.06),
                )
            except Exception:
                first.add_text(f"[图片无法插入：{Path(row[0]).name}]")
            paragraph.add_run("\t")
            second = paragraph.add_run()
            try:
                second.add_picture(
                    str(image_cache.word_image_path(row[1])),
                    width=Cm(7.21),
                    height=Cm(4.06),
                )
            except Exception:
                second.add_text(f"[图片无法插入：{Path(row[1]).name}]")


def _get_issue_value(issue: Any, *names: str) -> str:
    for name in names:
        if isinstance(issue, dict):
            value = issue.get(name)
        else:
            value = getattr(issue, name, None)
        if value:
            return str(value)
    return ""


def _get_issue_images(issue: Any) -> list[Path]:
    if isinstance(issue, dict):
        images = issue.get("image_paths") or issue.get("images") or []
    else:
        images = getattr(issue, "image_paths", None) or getattr(issue, "images", None) or []
    return [Path(image) for image in images if image]


def _get_issue_image_rows(issue: Any) -> list[list[Path]]:
    if isinstance(issue, dict):
        rows = issue.get("image_rows") or []
    else:
        rows = getattr(issue, "image_rows", None) or []
    normalized: list[list[Path]] = []
    for row in rows:
        normalized.append([Path(image) for image in row if image])
    return normalized


def _set_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    normal.font.size = Pt(10.5)

    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = FONT_HEITI if level == 1 else FONT_FANGSONG
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(14)
    run.bold = level in (1, 2)


def _add_body(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = FONT_FANGSONG
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FANGSONG)
    run.font.size = Pt(14)
    run.bold = bold

def _add_images(
    document: Document,
    images: list[Path],
    image_cache: ImageCache,
    width_cm: float = IMAGE_WIDTH_CM,
    height_cm: float = IMAGE_HEIGHT_CM,
) -> None:
    if not images:
        return

    for image in images:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        try:
            run.add_picture(str(image_cache.word_image_path(image)), width=Cm(width_cm), height=Cm(height_cm))
        except Exception:
            run2 = paragraph.add_run(f"[图片无法插入：{image.name}]")
            run2.font.size = Pt(9)

def _chinese_section_number(number: int) -> str:
    values = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五"}
    return values.get(number, str(number))



