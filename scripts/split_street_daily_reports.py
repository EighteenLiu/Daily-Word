from __future__ import annotations

import argparse
import hashlib
import re
import sys
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path


WORKSPACE = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parents[1]
INPUT_ROOT = WORKSPACE / "input"
OUTPUT_ROOT = WORKSPACE / "output"
REFERENCES_ROOT = INPUT_ROOT
XLSX_ROOT = OUTPUT_ROOT / "\u8f6c\u6362xlsx"
EXTRACT_DOCX_ROOT = OUTPUT_ROOT / "\u63d0\u53d6docx"
TRANSFER_ROOT = OUTPUT_ROOT / "\u4e2d\u8f6c\u7ad9\u65e5\u62a5"

CN_SONG = "\u5b8b\u4f53"
CN_HEITI = "\u9ed1\u4f53"
CN_FANGSONG = "\u4eff\u5b8b"
CN_STREET = "\u8857\u9053"
CN_REPORT_DIR = "\u8857\u9053\u65e5\u62a5"
CN_DISTRICT = "\u533a\u7ea7"
CN_CHECK_REPORT = "\u68c0\u67e5\u65e5\u62a5"
CN_NO_CONTENT = "\u65e0"
CN_PROBLEM_PREFIX = "\u5b58\u5728\u7684\u95ee\u9898\u662f\uff1a"
STYLE_HEADING_1 = "\u8bba\u65871"
STYLE_HEADING_2 = "\u8bba\u65871.1"
STYLE_HEADING_3 = "\u8bba\u65871.1.1"
STYLE_BODY = "\u5b8b5\u6b63\u6587"
PLACE_SUBHEADINGS = (
    "\u0031.\u5c0f\u533a\u6574\u4f53\u60c5\u51b5",
    "\u0032.\u6876\u7ad9\u8bbe\u7f6e\u60c5\u51b5",
    "\u0033.\u5c45\u6c11\u6295\u653e\u60c5\u51b5",
)
CN_TRANSFER_STATION = "\u4e2d\u8f6c\u7ad9"

TARGET_CATEGORIES = (
    "\u5c45\u4f4f\u5c0f\u533a\u3001\u5e73\u623f\u80e1\u540c",
    "\u9910\u996e\u5355\u4f4d",
    "\u793e\u4f1a\u5355\u4f4d",
)


@dataclass
class ImageBlob:
    blob: bytes
    content_type: str
    suffix: str
    name: str


@dataclass
class ContentBlock:
    text: str = ""
    images: list[ImageBlob] = field(default_factory=list)
    is_heading: bool = False
    heading_level: int | None = None


@dataclass
class PlaceBlock:
    name: str
    blocks: list[ContentBlock] = field(default_factory=list)


@dataclass
class TransferBlock:
    street: str
    title: str
    blocks: list[ContentBlock] = field(default_factory=list)


@dataclass
class ReportDate:
    value: date

    @property
    def short(self) -> str:
        return f"{self.value.month}.{self.value.day}"

    @property
    def chinese(self) -> str:
        return f"{self.value.year}\u5e74{self.value.month}\u6708{self.value.day}\u65e5"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Split the generated structured docx into one street daily report per street. "
            "Defaults are normalized for this project: template from input, source from output/提取docx, "
            "and reports written to output/street-daily-report(date)."
        )
    )
    parser.add_argument(
        "input",
        nargs="?",
        type=Path,
        help="Structured .docx file. Defaults to the newest structured .docx in output/提取docx/.",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        help="Output directory. Defaults to output/街道日报（日期）.",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=REFERENCES_ROOT / "\u65e5\u62a5\u6a21\u7248.docx",
        help="Daily report template docx. Defaults to input/日报模版.docx.",
    )
    parser.add_argument(
        "--date",
        help="Report date such as 5.17, 2026-05-17, or 20260517. Auto-detected when omitted.",
    )
    parser.add_argument(
        "--source-xlsx",
        type=Path,
        help="Converted xlsx used only for date inference. Defaults to the xlsx matching the input docx prefix.",
    )
    parser.add_argument(
        "--transfer-doc",
        type=Path,
        help="Optional transfer-station daily report .doc/.docx. Defaults to matching date file in input/ or output/中转站日报/.",
    )
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing report files.")
    parser.add_argument(
        "--include-non-street-headings",
        action="store_true",
        help="Also generate reports for Heading 2 names that do not contain '街道'.",
    )
    return parser.parse_args()


def strip_heading_number(text: str) -> str:
    text = text.strip()
    patterns = (
        r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u96f6]+[\u3001.．]\s*",
        r"^[（(][\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u96f6]+[）)]\s*",
        r"^\d+[.．、]\s*",
        r"^[（(]\d+[）)]\s*",
    )
    for pattern in patterns:
        text = re.sub(pattern, "", text)
    return text.strip()


def chinese_numeral(number: int) -> str:
    digits = "\u96f6\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d"
    units = ["", "\u5341", "\u767e", "\u5343"]
    if number <= 0:
        return str(number)
    if number < 10:
        return digits[number]
    if number < 20:
        return "\u5341" + (digits[number % 10] if number % 10 else "")

    chars: list[str] = []
    zero_pending = False
    text = str(number)
    length = len(text)
    for index, char in enumerate(text):
        digit = int(char)
        unit_index = length - index - 1
        if digit == 0:
            zero_pending = bool(chars)
            continue
        if zero_pending:
            chars.append(digits[0])
            zero_pending = False
        chars.append(digits[digit] + units[unit_index])
    return "".join(chars)


def safe_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\r\n\t]+', "_", name.strip())
    name = re.sub(r"\s+", " ", name).strip(" .")
    return name or "report"


def short_street_name(street: str) -> str:
    return street[:-2] if street.endswith(CN_STREET) else street


def find_default_input() -> Path:
    candidates = [
        path
        for path in EXTRACT_DOCX_ROOT.glob("*.docx")
        if not path.name.startswith("~$")
        and "\u7ed3\u6784\u5316\u63d0\u53d6" in path.name
    ]
    if not candidates:
        candidates = [
            path
            for path in OUTPUT_ROOT.glob("*.docx")
            if not path.name.startswith("~$")
            and "\u7ed3\u6784\u5316\u63d0\u53d6" in path.name
        ]
    if not candidates:
        raise FileNotFoundError("No structured extraction .docx found in output/提取docx/.")
    return max(candidates, key=lambda path: path.stat().st_mtime)


def parse_report_date(value: str) -> ReportDate:
    value = value.strip()
    if re.fullmatch(r"\d{8}", value):
        return ReportDate(datetime.strptime(value, "%Y%m%d").date())
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", value):
        return ReportDate(datetime.strptime(value, "%Y-%m-%d").date())
    if re.fullmatch(r"\d{1,2}\.\d{1,2}", value):
        month, day = value.split(".")
        return ReportDate(date(datetime.now().year, int(month), int(day)))
    raise ValueError(f"Unsupported date format: {value}")


def matching_xlsx_for_docx(input_path: Path) -> Path | None:
    stem = input_path.stem.split("_\u7ed3\u6784\u5316\u63d0\u53d6", 1)[0]
    candidate = XLSX_ROOT / f"{stem}.xlsx"
    if candidate.exists():
        return candidate
    matches = sorted(XLSX_ROOT.glob("*.xlsx"), key=lambda path: path.stat().st_mtime, reverse=True)
    if not matches:
        matches = sorted(OUTPUT_ROOT.glob("*.xlsx"), key=lambda path: path.stat().st_mtime, reverse=True)
    return matches[0] if matches else None


def infer_date_from_xlsx(xlsx_path: Path) -> ReportDate | None:
    try:
        import openpyxl
    except ImportError:
        return None

    workbook = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    try:
        sheet = workbook[workbook.sheetnames[0]]
        for row in range(1, min(sheet.max_row, 30) + 1):
            for col in range(1, min(sheet.max_column, 10) + 1):
                value = sheet.cell(row, col).value
                match = re.search(r"(20\d{6})", str(value or ""))
                if match:
                    return parse_report_date(match.group(1))
    finally:
        workbook.close()
    return None


def infer_latest_date_from_xlsx(xlsx_path: Path) -> ReportDate | None:
    try:
        import openpyxl
    except ImportError:
        return None

    dates: list[date] = []
    workbook = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    try:
        sheet = workbook[workbook.sheetnames[0]]
        identifier_columns = identifier_date_columns(sheet)
        if not identifier_columns:
            return None
        for row_index, row in enumerate(sheet.iter_rows(), start=1):
            for column_index, cell in enumerate(row, start=1):
                if column_index not in identifier_columns:
                    continue
                dates.extend(dates_from_identifier_value(cell.value))
    finally:
        workbook.close()
    return ReportDate(max(dates)) if dates else None


def identifier_date_columns(sheet) -> set[int]:
    columns: set[int] = set()
    max_row = min(sheet.max_row, 3)
    for row in range(1, max_row + 1):
        for col in range(1, sheet.max_column + 1):
            header = re.sub(r"\s+", "", str(sheet.cell(row, col).value or ""))
            if "编号" in header:
                columns.add(col)
    return columns


def dates_from_identifier_value(value: object) -> list[date]:
    if value is None:
        return []
    text = str(value).strip()
    if not text:
        return []
    digits = re.sub(r"\D", "", text)
    if len(digits) < 8:
        return []
    match = re.search(r"20\d{6}", digits)
    if not match:
        return []
    parsed = _safe_date(match.group(0)[:4], match.group(0)[4:6], match.group(0)[6:8])
    return [parsed] if parsed else []


def dates_from_cell_value(value: object, allow_short_dates: bool = True) -> list[date]:
    if value is None:
        return []
    if isinstance(value, datetime):
        return [value.date()]
    if isinstance(value, date):
        return [value]

    text = str(value)
    dates: list[date] = []
    patterns = (
        r"(20\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})",
        r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)",
        r"(?<!\d)(20\d{2})(\d{2})(\d{2})\d+",
    )
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            parsed = _safe_date(match.group(1), match.group(2), match.group(3))
            if parsed:
                dates.append(parsed)

    current_year = datetime.now().year
    if allow_short_dates:
        for match in re.finditer(r"(?<!\d)(\d{1,2})月(\d{1,2})日", text):
            parsed = _safe_date(str(current_year), match.group(1), match.group(2))
            if parsed:
                dates.append(parsed)
        for match in re.finditer(r"(?<!\d)(\d{1,2})[.．](\d{1,2})(?!\d)", text):
            parsed = _safe_date(str(current_year), match.group(1), match.group(2))
            if parsed:
                dates.append(parsed)
    return dates


def _safe_date(year: str, month: str, day: str) -> date | None:
    try:
        return date(int(year), int(month), int(day))
    except ValueError:
        return None


def infer_date_from_filename(path: Path) -> ReportDate | None:
    match = re.search(r"(20\d{2})[-_]?(\d{2})[-_]?(\d{2})", path.name)
    if not match:
        return None
    detected = date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    return ReportDate(detected)


def resolve_report_date(args: argparse.Namespace, input_path: Path) -> ReportDate:
    if args.date:
        return parse_report_date(args.date)

    source_xlsx = args.source_xlsx.resolve() if args.source_xlsx else matching_xlsx_for_docx(input_path)
    if source_xlsx and source_xlsx.exists():
        inferred = infer_latest_date_from_xlsx(source_xlsx)
        if inferred:
            return inferred

    raise RuntimeError("Could not infer report date from the 编号 column. Pass --date 5.17 or --date 2026-05-17.")


def transfer_date_tokens(report_date: ReportDate) -> tuple[str, ...]:
    return (
        f"{report_date.value.month}\u6708{report_date.value.day}\u65e5",
        f"{report_date.value.month}.{report_date.value.day}",
        f"{report_date.value.month}-{report_date.value.day}",
        f"{report_date.value.month}_{report_date.value.day}",
        report_date.value.strftime("%Y%m%d"),
        report_date.value.strftime("%Y-%m-%d"),
    )


def transfer_doc_matches_report_date(path: Path, report_date: ReportDate) -> bool:
    return any(token in path.name for token in transfer_date_tokens(report_date))


def find_transfer_doc(report_date: ReportDate, explicit: Path | None = None) -> Path | None:
    if explicit:
        path = explicit.resolve()
        if not path.exists():
            return None
        if not transfer_doc_matches_report_date(path, report_date):
            print(
                f"[skip] transfer_doc date mismatch: {path.name} "
                f"(expected {report_date.value.month}\u6708{report_date.value.day}\u65e5)"
            )
            return None
        return path

    date_tokens = transfer_date_tokens(report_date)
    search_roots = [
        INPUT_ROOT,
        INPUT_ROOT / "references",
        OUTPUT_ROOT / "references",
        TRANSFER_ROOT,
    ]
    candidates: list[Path] = []
    for root in search_roots:
        if not root.exists():
            continue
        for suffix in ("*.docx", "*.doc"):
            for path in root.glob(suffix):
                if path.name.startswith("~$"):
                    continue
                if CN_TRANSFER_STATION in path.name and any(token in path.name for token in date_tokens):
                    candidates.append(path)
    if candidates:
        return max(candidates, key=lambda path: path.stat().st_mtime)
    return None


def convert_doc_to_docx(source: Path) -> Path:
    if source.suffix.lower() == ".docx":
        return source

    try:
        import win32com.client as win32
    except ImportError as exc:
        raise RuntimeError("pywin32 is required to convert transfer-station .doc files.") from exc

    TRANSFER_ROOT.mkdir(parents=True, exist_ok=True)
    destination = TRANSFER_ROOT / f"{source.stem}.docx"
    if destination.exists() and destination.stat().st_mtime >= source.stat().st_mtime:
        return destination

    word = None
    document = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(source.resolve()), ReadOnly=True, AddToRecentFiles=False)
        document.SaveAs2(str(destination.resolve()), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
    return destination


def convert_template_to_docx(source: Path) -> Path:
    if source.suffix.lower() == ".docx":
        return source

    try:
        import win32com.client as win32
    except ImportError as exc:
        raise RuntimeError("pywin32 is required to convert .doc templates.") from exc

    destination = source.with_suffix(".docx")
    if destination.exists() and destination.stat().st_mtime >= source.stat().st_mtime:
        return destination

    word = None
    document = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(source.resolve()), ReadOnly=True, AddToRecentFiles=False)
        document.SaveAs2(str(destination.resolve()), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
    return destination


def street_from_transfer_title(title: str, known_streets: set[str]) -> str | None:
    for street in sorted(known_streets, key=len, reverse=True):
        if title.startswith(street) or street in title:
            return street
    return None


def parse_transfer_docx(path: Path, known_streets: set[str]) -> dict[str, list[TransferBlock]]:
    from docx import Document

    if not path or not path.exists():
        return {}

    docx_path = convert_doc_to_docx(path)
    document = Document(docx_path)
    transfers: dict[str, list[TransferBlock]] = {}
    current: TransferBlock | None = None

    for paragraph in document.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()
        images = extract_paragraph_images(paragraph)

        if style.startswith("Heading") and text:
            street = street_from_transfer_title(text, known_streets)
            if street:
                current = TransferBlock(street=street, title=text)
                transfers.setdefault(street, []).append(current)
            else:
                current = None
            continue

        if current is None:
            continue
        if not text and not images:
            continue
        is_transfer_heading = bool(re.fullmatch(r"\d+[.．、]\s*(整体情况|具体情况)", text))
        is_transfer_subheading = bool(re.fullmatch(r"[（(]\d+[）)].+", text))
        heading_level = 3 if is_transfer_heading else 4 if is_transfer_subheading else None
        current.blocks.append(
            ContentBlock(
                text=text,
                images=images,
                is_heading=heading_level is not None,
                heading_level=heading_level,
            )
        )

    return transfers


def extract_paragraph_images(paragraph) -> list[ImageBlob]:
    from docx.oxml.ns import qn

    images: list[ImageBlob] = []
    image_nodes = list(paragraph._p.xpath('.//*[local-name()="blip"]'))
    image_nodes.extend(paragraph._p.xpath('.//*[local-name()="imagedata"]'))
    for node in image_nodes:
        r_id = node.get(qn("r:embed")) or node.get(qn("r:id"))
        if not r_id:
            continue
        image_part = paragraph.part.related_parts.get(r_id)
        if image_part is None:
            continue
        suffix = "." + image_part.partname.ext
        images.append(
            ImageBlob(
                blob=image_part.blob,
                content_type=image_part.content_type,
                suffix=suffix,
                name=Path(str(image_part.partname)).name,
            )
        )
    return images


def parse_structured_docx(
    input_path: Path,
    include_non_street_headings: bool,
) -> dict[str, dict[str, dict[str, PlaceBlock]]]:
    from docx import Document

    document = Document(input_path)
    reports: dict[str, dict[str, dict[str, PlaceBlock]]] = {}
    current_category = ""
    current_street = ""
    current_place = ""

    for paragraph in document.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()

        if style == "Heading 1":
            current_category = strip_heading_number(text)
            current_street = ""
            current_place = ""
            continue

        if style == "Heading 2":
            current_street = strip_heading_number(text)
            current_place = ""
            if not include_non_street_headings and CN_STREET not in current_street:
                current_street = ""
                continue
            if current_category in TARGET_CATEGORIES and current_street:
                reports.setdefault(current_street, {}).setdefault(current_category, {})
            continue

        if style == "Heading 3":
            current_place = strip_heading_number(text)
            if current_category in TARGET_CATEGORIES and current_street and current_place:
                category = reports.setdefault(current_street, {}).setdefault(current_category, {})
                category.setdefault(current_place, PlaceBlock(current_place))
            continue

        if current_category not in TARGET_CATEGORIES or not current_street or not current_place:
            continue

        images = extract_paragraph_images(paragraph)
        if not text and not images:
            continue

        category = reports.setdefault(current_street, {}).setdefault(current_category, {})
        place = category.setdefault(current_place, PlaceBlock(current_place))
        place.blocks.append(ContentBlock(text=text, images=images))

    return reports


def remove_template_body_content(document) -> None:
    body = document._body._element
    for element in list(body):
        if element.tag.endswith("sectPr"):
            continue
        if element.tag.endswith("tbl"):
            continue
        body.remove(element)


def update_template_date(document, report_date: ReportDate) -> None:
    date_pattern = re.compile(r"\d{4}\u5e74\d{1,2}\u6708\d{1,2}\u65e5")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = date_pattern.sub(report_date.chinese, run.text)


def set_document_styles(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = CN_SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_SONG)
    normal.font.size = Pt(10.5)


def new_document(template_path: Path, report_date: ReportDate):
    from docx import Document

    if template_path.exists():
        template_path = convert_template_to_docx(template_path)
        document = Document(template_path)
        update_template_date(document, report_date)
        remove_template_body_content(document)
    else:
        document = Document()
    set_document_styles(document)
    return document


def add_raw_picture_from_blob(run, image: ImageBlob, width: int, height: int | None = None) -> None:
    from PIL import Image
    from docx.oxml.shape import CT_Inline
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.parts.image import ImagePart
    from docx.shared import Emu

    try:
        with Image.open(BytesIO(image.blob)) as pil_image:
            pixel_width, pixel_height = pil_image.size
    except Exception as exc:
        raise RuntimeError(f"Cannot read image dimensions without recompression: {image.name}") from exc

    image_parts = run.part.package.image_parts
    digest = hashlib.sha1(image.blob).hexdigest()
    image_part = image_parts._get_by_sha1(digest)
    if image_part is None:
        partname = image_parts._next_image_partname(image.suffix.lstrip("."))
        image_part = ImagePart(partname, image.content_type, image.blob)
        image_parts.append(image_part)

    r_id = run.part.relate_to(image_part, RT.IMAGE)
    display_height = Emu(height if height is not None else int(width * pixel_height / pixel_width))
    inline = CT_Inline.new_pic_inline(run.part.next_id, r_id, image.name, width, display_height)
    run._r.add_drawing(inline)


def image_display_size(image: ImageBlob) -> tuple[int, int]:
    from PIL import Image, ExifTags
    from docx.shared import Cm

    landscape = (int(Cm(10.17)), int(Cm(5.72)))
    portrait = (int(Cm(5.72)), int(Cm(10.17)))
    square = (int(Cm(8.0)), int(Cm(8.0)))

    try:
        with Image.open(BytesIO(image.blob)) as pil_image:
            width, height = pil_image.size
            orientation = None
            exif = pil_image.getexif()
            if exif:
                orientation_tag = next(
                    (key for key, value in ExifTags.TAGS.items() if value == "Orientation"),
                    None,
                )
                orientation = exif.get(orientation_tag) if orientation_tag else None
            if orientation in (5, 6, 7, 8):
                width, height = height, width
    except Exception:
        return landscape

    if width > height * 1.1:
        return landscape
    if height > width * 1.1:
        return portrait
    return square


def add_images(document, images: list[ImageBlob]) -> None:
    if not images:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Emu, Pt
    from PIL import Image

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    for idx, image in enumerate(images):
        with Image.open(BytesIO(image.blob)) as pil_img:
            w, h = pil_img.size
        img_width = int(Cm(10.2))
        img_height = Emu(int(int(img_width) * h / w))
        add_raw_picture_from_blob(run, image, img_width, img_height)
        if idx < len(images) - 1:
            run.add_text("  ")

def add_section_heading(document, text: str, level: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style_names = {1: STYLE_HEADING_1, 2: STYLE_HEADING_2, 3: STYLE_HEADING_3}
    paragraph = document.add_paragraph()
    style_name = style_names.get(level)
    if style_name and style_name in [style.name for style in document.styles]:
        paragraph.style = style_name
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(min(max(level, 1), 4) - 1))
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(3)
    if level == 2:
        paragraph.paragraph_format.left_indent = Pt(10.5)
    elif level == 3:
        paragraph.paragraph_format.left_indent = Pt(21)
    elif level >= 4:
        paragraph.paragraph_format.left_indent = Pt(31.5)
    run = paragraph.add_run(text)
    run.bold = level in (1, 4)
    font_name = CN_HEITI if level == 1 else CN_FANGSONG
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(14)


def add_problem_paragraph(document, text: str) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph(text)
    if STYLE_BODY in [style.name for style in document.styles]:
        paragraph.style = STYLE_BODY
    for run in paragraph.runs:
        run.font.name = CN_FANGSONG
        run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FANGSONG)
        run.font.size = Pt(14)


def add_problem_summary_paragraph(document, texts: list[str]) -> None:
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.shared import Pt

    merged_text = "\uff1b".join(text.strip() for text in texts if text.strip())
    paragraph = document.add_paragraph()
    if STYLE_BODY in [style.name for style in document.styles]:
        paragraph.style = STYLE_BODY
    prefix_run = paragraph.add_run(CN_PROBLEM_PREFIX + " ")
    prefix_run.font.name = CN_FANGSONG
    prefix_run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FANGSONG)
    prefix_run.font.size = Pt(14)
    if merged_text:
        problem_run = paragraph.add_run(merged_text)
        problem_run.font.name = CN_FANGSONG
        problem_run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FANGSONG)
        problem_run.font.size = Pt(14)
        problem_run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_body_paragraph(document, text: str) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    if STYLE_BODY in [style.name for style in document.styles]:
        paragraph.style = STYLE_BODY
    run = paragraph.add_run(text)
    run.font.name = CN_FANGSONG
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FANGSONG)
    run.font.size = Pt(14)


def add_place_overall_content(document, blocks: list[ContentBlock]) -> None:
    import re

    CN_PROMO = "\u5c0f\u533a\u5ba3\u4f20\u6c1b\u56f4"
    CN_NOTICE = "\u5c0f\u533a\u516c\u793a\u724c"
    CN_ATMOSPHERE = "\u5ba3\u4f20\u6c1b\u56f4"
    NO_PROBLEM_PREFIXES = (
        "\u65e0\u95ee\u9898",
        "\u65e0",
        "\u672a\u53d1\u73b0\u95ee\u9898",
        "\u672a\u89c1\u95ee\u9898",
        "\u672a\u89c1\u660e\u663e\u95ee\u9898",
    )

    problem_texts: list[str] = []
    label_blocks: list[tuple[str, list[ImageBlob]]] = []
    problem_images: list[ImageBlob] = []

    for block in blocks:
        text = block.text.strip() if block.text else ""
        cleaned = re.sub(r"^[\uff08(]\d+[\uff09)]\s*", "", text) if text else ""

        if CN_PROMO in cleaned or CN_NOTICE in cleaned or CN_ATMOSPHERE in cleaned:
            label_blocks.append((cleaned, block.images))
        elif cleaned:
            problem_texts.append(cleaned)
            if block.images:
                problem_images.extend(block.images)
        else:
            # Image-only block (no meaningful text)
            if block.images:
                problem_images.extend(block.images)

    real_problems = [
        t
        for t in problem_texts
        if t and not any(t.strip().startswith(prefix) for prefix in NO_PROBLEM_PREFIXES)
    ]
    if real_problems:
        add_problem_summary_paragraph(document, problem_texts)
    elif problem_texts:
        add_problem_summary_paragraph(document, ["\u65e0\u95ee\u9898"])
    elif not label_blocks:
        add_problem_summary_paragraph(document, ["\u65e0\u95ee\u9898"])

    if problem_images:
        add_images(document, problem_images)

    for label_text, images in label_blocks:
        add_body_paragraph(document, label_text)
        if images:
            add_images(document, images)

def write_report(
    output_path: Path,
    template_path: Path,
    report_date: ReportDate,
    street: str,
    categories: dict[str, dict[str, PlaceBlock]],
    transfer_blocks: list[TransferBlock] | None = None,
) -> None:
    document = new_document(template_path, report_date)

    visible_category_index = 0
    for category_name in TARGET_CATEGORIES:
        places = categories.get(category_name, {})
        if not places:
            continue

        visible_category_index += 1
        add_section_heading(document, f"{chinese_numeral(visible_category_index)}\u3001{category_name}", level=1)

        for place_index, place in enumerate(places.values(), start=1):
            add_section_heading(document, f"\uff08{chinese_numeral(place_index)}\uff09{place.name}", level=2)
            for subheading_index, subheading in enumerate(PLACE_SUBHEADINGS):
                add_section_heading(document, subheading, level=3)
                if subheading_index != 0:
                    continue
                add_place_overall_content(document, place.blocks)

    if transfer_blocks:
        visible_category_index += 1
        add_section_heading(document, f"{chinese_numeral(visible_category_index)}\u3001{CN_TRANSFER_STATION}", level=1)
        for transfer_index, transfer in enumerate(transfer_blocks, start=1):
            add_section_heading(document, f"\uff08{chinese_numeral(transfer_index)}\uff09{transfer.title}", level=2)
            transfer_images: list[ImageBlob] = []
            for block in transfer.blocks:
                if block.text:
                    if block.is_heading:
                        add_section_heading(document, block.text, level=block.heading_level or 3)
                    else:
                        add_problem_paragraph(document, block.text)
                transfer_images.extend(block.images)
            add_images(document, transfer_images)

    if visible_category_index == 0:
        document.add_paragraph(CN_NO_CONTENT)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    return run_split(args)


def run_split(args: argparse.Namespace) -> int:
    try:
        input_path = args.input.resolve() if args.input else find_default_input()
        if not input_path.exists():
            raise FileNotFoundError(f"Input file does not exist: {input_path}")

        report_date = resolve_report_date(args, input_path)
        output_dir = (
            args.output_dir.resolve()
            if args.output_dir
            else OUTPUT_ROOT / f"{CN_REPORT_DIR}\uff08{report_date.short}\uff09"
        )
        template_path = args.template.resolve()
        reports = parse_structured_docx(input_path, include_non_street_headings=args.include_non_street_headings)
        if not reports:
            raise RuntimeError("No street content found. Check Heading 1/2/3 structure.")
        transfer_doc = find_transfer_doc(report_date, explicit=args.transfer_doc)
        transfers = parse_transfer_docx(transfer_doc, set(reports.keys())) if transfer_doc else {}
        if transfer_doc:
            print(f"[ok] transfer_doc: {transfer_doc}")

        written = 0
        for street, categories in reports.items():
            filename = f"{report_date.short}{CN_DISTRICT}{short_street_name(street)}{CN_CHECK_REPORT}.docx"
            output_path = output_dir / safe_filename(filename)
            if output_path.exists() and not args.overwrite:
                print(f"[skip] exists: {output_path}")
                continue
            write_report(output_path, template_path, report_date, street, categories, transfers.get(street))
            written += 1
            print(f"[ok] {street} -> {output_path}")

        print(f"[ok] date: {report_date.short}")
        print(f"[ok] output_dir: {output_dir}")
        print(f"[ok] streets: {len(reports)}, written: {written}")
        return 0
    except Exception as exc:
        print(f"[fail] {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())


