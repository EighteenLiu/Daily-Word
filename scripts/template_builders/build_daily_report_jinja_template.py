from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "generated" / "daily_report_jinja_template.docx"

FONT_SONG = "宋体"
FONT_FANGSONG = "仿宋"
FONT_HEITI = "黑体"


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_para(document: Document, text: str = "", font: str = FONT_FANGSONG, size: int = 14, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return paragraph


def add_heading(document: Document, text: str, level: int) -> None:
    font = FONT_HEITI if level == 1 else FONT_FANGSONG
    bold = level in (1, 2)
    paragraph = add_para(document, text, font=font, size=14, bold=bold)
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(3)


def add_jinja(document: Document, tag: str) -> None:
    add_para(document, tag, font="Courier New", size=10)


def add_image_loop(document: Document, loop_var: str, image_var: str = "image") -> None:
    add_jinja(document, f"{{%p for {image_var} in {loop_var} %}}")
    paragraph = add_para(document, f"{{{{ {image_var} }}}}", font=FONT_FANGSONG, size=12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_jinja(document, "{%p endfor %}")


def section_title(name: str) -> str:
    return "{{ ['','一','二','三','四','五','六','七','八','九','十'][section_no.n] }}、" + name


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = document.styles["Normal"]
    normal.font.name = FONT_SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    normal.font.size = Pt(10.5)

    title = add_para(document, "{{ report_title | default('区级检查日报') }}", font=FONT_HEITI, size=16, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(document, "{{ report_date_text }}", font=FONT_SONG, size=12)

    add_jinja(document, "{%p set section_no = namespace(n=0) %}")
    add_jinja(document, "{%p if communities %}")
    add_jinja(document, "{%p set section_no.n = section_no.n + 1 %}")
    add_heading(document, section_title("居住小区"), level=1)
    add_jinja(document, "{%p for community in communities %}")
    add_heading(document, "（{{ community.index_cn }}）{{ community.name }}", level=2)
    add_heading(document, "1.小区整体情况", level=3)
    add_para(document, "{{ community.overall_intro }}存在的问题是：{{ community.overall_problem_summary }}")
    add_para(document, "（1）小区宣传氛围：{% if community.promo_text != \"无问题\" %}{{ community.promo_text }}{% endif %}")
    add_image_loop(document, "community.promo_images")
    add_para(document, "（2）小区公示牌：")
    add_image_loop(document, "community.notice_board_images")
    add_jinja(document, "{%p if community.is_pure_box_room %}")
    add_para(document, "（3）装修垃圾投放点设置")
    add_para(document, "预约收集，集中密闭运输。")
    add_para(document, "（4）大件垃圾投放点设置")
    add_para(document, "预约收集，集中密闭运输。")
    add_jinja(document, "{%p endif %}")
    add_heading(document, "2.桶站设置情况", level=3)
    add_jinja(document, "{%p for station in community.stations %}")
    add_para(document, "{{ station.station_no }}号桶站设置情况：{{ station.problem_summary }}")
    add_image_loop(document, "station.images")
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p if community.resident_delivery %}")
    add_heading(document, "3.居民投放情况：{{ community.resident_delivery.summary }}", level=3)
    add_image_loop(document, "community.resident_delivery.error_images")
    add_jinja(document, "{%p endif %}")
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p endif %}")

    add_jinja(document, "{%p if restaurants %}")
    add_jinja(document, "{%p set section_no.n = section_no.n + 1 %}")
    add_heading(document, section_title("餐饮单位"), level=1)
    add_jinja(document, "{%p for restaurant in restaurants %}")
    add_heading(document, "（{{ restaurant.index_cn }}）{{ restaurant.name }}", level=2)
    add_heading(document, "1.整体情况", level=3)
    add_para(document, "存在的问题是：{{ restaurant.overall_problem_summary }}")
    add_para(document, "（1）宣传氛围：{% if restaurant.promo_text != \"无问题\" %}{{ restaurant.promo_text }}{% endif %}")
    add_image_loop(document, "restaurant.promo_images")
    add_heading(document, "2.桶站设置情况：{{ restaurant.container_problem_summary }}", level=3)
    add_image_loop(document, "restaurant.container_images")
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p endif %}")

    add_jinja(document, "{%p if social_units %}")
    add_jinja(document, "{%p set section_no.n = section_no.n + 1 %}")
    add_heading(document, section_title("社会单位"), level=1)
    add_jinja(document, "{%p for social_unit in social_units %}")
    add_heading(document, "（{{ social_unit.index_cn }}）{{ social_unit.name }}", level=2)
    add_heading(document, "1.整体情况", level=3)
    add_para(document, "存在的问题是：{{ social_unit.overall_problem_summary }}")
    add_para(document, "（1）宣传氛围：{% if social_unit.promo_text != \"无问题\" %}{{ social_unit.promo_text }}{% endif %}")
    add_image_loop(document, "social_unit.promo_images")
    add_heading(document, "2.桶站设置情况：{{ social_unit.container_problem_summary }}", level=3)
    add_image_loop(document, "social_unit.container_images")
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p endif %}")

    add_jinja(document, "{%p if outside_bucket_issues %}")
    add_jinja(document, "{%p set section_no.n = section_no.n + 1 %}")
    add_heading(document, section_title("桶外摆检查"), level=1)
    add_jinja(document, "{%p for issue in outside_bucket_issues %}")
    add_para(document, "{{ issue.street_name }}：{{ issue.clean_text }}")
    add_jinja(document, "{%p for row in issue.image_rows %}")
    paragraph = add_para(
        document,
        "{{ row[0] }}{% if row|length > 1 %}\t{{ row[1] }}{% endif %}",
        font=FONT_FANGSONG,
        size=12,
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p endfor %}")
    add_jinja(document, "{%p endif %}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
