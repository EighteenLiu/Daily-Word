from __future__ import annotations

import re
import sys
from collections import OrderedDict
from datetime import timedelta
from pathlib import Path
from typing import Iterable


WORKSPACE = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parents[1]
OUTPUT_ROOT = WORKSPACE / "output"
SUMMARY_ROOT = OUTPUT_ROOT / "\u6bcf\u65e5\u603b\u7ed3"

CN_COMMUNITY_CATEGORY = "\u5c45\u4f4f\u5c0f\u533a\u3001\u5e73\u623f\u80e1\u540c"
CN_PROBLEM_PREFIX = "\u5b58\u5728\u7684\u95ee\u9898\u662f\uff1a"
CN_DISTRICT = "\u533a\u7ea7"
CN_GARBAGE_DAILY = "\u5783\u573e\u5206\u7c7b\u5de5\u4f5c\u65e5\u62a5"
CN_DAILY_SUMMARY = "\u6bcf\u65e5\u6c47\u603b\u60c5\u51b5"
CN_STREET = "\u8857\u9053"
FONT_FANGSONG_GB2312 = "\u4eff\u5b8b_GB2312"
FONT_SONG = "\u5b8b\u4f53"


def chinese_numeral(number: int) -> str:
    try:
        import split_street_daily_reports as split
    except ModuleNotFoundError:
        from scripts import split_street_daily_reports as split

    return split.chinese_numeral(number)


def short_street_name(street: str) -> str:
    return street[:-2] if street.endswith(CN_STREET) else street


def full_street_name(short_name: str) -> str:
    return short_name if short_name.endswith(CN_STREET) else f"{short_name}{CN_STREET}"


def clean_problem_number(text: str) -> str:
    text = text.strip()
    while True:
        cleaned = re.sub(r"^[\uff08(]\d+[\uff09)]\s*", "", text)
        if cleaned == text:
            break
        text = cleaned.strip()
    return text.strip()


def split_problem_segments(text: str) -> list[str]:
    parts = re.split(r"(?=[\uff08(]\d+[\uff09)])", text.strip())
    return [part.strip().strip("\uff1b;") for part in parts if part.strip().strip("\uff1b;")]


def is_real_problem(text: str) -> bool:
    cleaned = clean_problem_number(text)
    if not cleaned:
        return False
    no_problem_prefixes = (
        "\u65e0\u95ee\u9898",
        "\u672a\u53d1\u73b0\u95ee\u9898",
        "\u672a\u89c1\u95ee\u9898",
        "\u65e0\u660e\u663e\u95ee\u9898",
    )
    return not cleaned.startswith(no_problem_prefixes)


def is_no_problem_summary(text: str) -> bool:
    cleaned = clean_problem_number(text).strip("：:，,。；;、 ")
    if not cleaned:
        return True
    no_problem_values = {
        "无",
        "无问题",
        "没问题",
        "未发现问题",
        "未见问题",
        "无明显问题",
    }
    return cleaned in no_problem_values


def clean_report_problem_item(text: str) -> str:
    text = clean_problem_number(text)
    text = text.strip("：:，,。；;、 ")
    text = clean_problem_number(text)
    text = text.strip("：:，,。；;、 ")
    return text


def place_problem_texts(place) -> list[str]:
    texts: list[str] = []
    for block in place.blocks:
        if block.text and block.text.strip():
            texts.append(block.text.strip())
    return texts


def merge_problem_texts(texts: list[str]) -> str:
    return "\uff1b".join(text.strip().strip("\uff1b;") for text in texts if text.strip())


def count_real_problems(texts: list[str]) -> int:
    count = 0
    for text in texts:
        for segment in split_problem_segments(text):
            if is_real_problem(segment):
                count += problem_segment_count(segment)
    return count


def problem_segment_count(text: str) -> int:
    cleaned = clean_problem_number(text)
    explicit_counts = [int(match) for match in re.findall(r"(\d+)\s*处", cleaned)]
    explicit_counts.extend(_parse_chinese_count(value) for value in re.findall(r"([零〇一二两三四五六七八九十百]+)\s*处", cleaned))
    explicit_counts = [value for value in explicit_counts if value is not None]
    return max(explicit_counts) if explicit_counts else 1


def _parse_chinese_count(value: str) -> int | None:
    if value.isdigit():
        return int(value)
    mapping = {
        "零": 0,
        "〇": 0,
        "一": 1,
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    if value in mapping:
        return mapping[value]
    if "百" in value:
        left, right = value.split("百", 1)
        hundred = mapping.get(left, 1 if not left else 0) * 100
        return hundred + (_parse_chinese_count(right) or 0)
    if value.startswith("十"):
        return 10 + (_parse_chinese_count(value[1:]) or 0)
    if "十" in value:
        left, right = value.split("十", 1)
        return mapping.get(left, 0) * 10 + (_parse_chinese_count(right) or 0)
    return None


def street_report_problem_map(street_report_paths: Iterable[tuple[str, Path]]) -> "OrderedDict[str, OrderedDict[str, str]]":
    result: OrderedDict[str, OrderedDict[str, str]] = OrderedDict()
    for street, path in street_report_paths:
        place_problems = extract_garbage_problem_texts_from_street_report(path)
        if place_problems:
            result[street] = place_problems
    return result


def collect_street_report_paths(report_dir: Path) -> list[tuple[str, Path]]:
    report_paths: list[tuple[str, Path]] = []
    for path in sorted(report_dir.glob("*.docx"), key=lambda item: item.name):
        if path.name.startswith("~$"):
            continue
        street = _street_name_from_report_filename(path)
        if street:
            report_paths.append((street, path))
    if not report_paths:
        raise RuntimeError(f"未在文件夹中找到街道日报 docx：{report_dir}")
    return report_paths


def infer_report_date_from_street_reports(report_paths: Iterable[tuple[str, Path]]):
    try:
        import split_street_daily_reports as split
    except ModuleNotFoundError:
        from scripts import split_street_daily_reports as split

    dates = []
    for _, path in report_paths:
        match = re.search(r"(\d{1,2}\.\d{1,2})", path.name)
        if match:
            dates.append(split.parse_report_date(match.group(1)))
            continue
        inferred = split.infer_date_from_filename(path)
        if inferred:
            dates.append(inferred)
    if not dates:
        raise RuntimeError("无法从已有日报文件名识别日期，请确认文件名包含类似 6.16 或 20260616 的日期。")
    return max(dates, key=lambda item: item.value)


def _street_name_from_report_filename(path: Path) -> str:
    text = path.stem
    match = re.search(r"区级(.+?)检查日报", text)
    if match:
        return full_street_name(match.group(1).strip())
    match = re.search(r"(.+?)街道", text)
    if match:
        return f"{match.group(1).strip()}街道"
    return ""


def extract_garbage_problem_texts_from_street_report(path: Path) -> "OrderedDict[str, str]":
    from docx import Document

    document = Document(path)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    communities: OrderedDict[str, list[str]] = OrderedDict()
    current_place: str | None = None
    in_community_section = False

    for text in texts:
        if text.startswith("一、居住小区"):
            in_community_section = True
            continue
        if in_community_section and re.match(r"^[二三四五六七八九十]、", text):
            break
        if not in_community_section:
            continue

        place_match = re.fullmatch(r"（[一二三四五六七八九十]+）(.+)", text)
        if place_match:
            current_place = place_match.group(1).strip()
            communities.setdefault(current_place, [])
            continue
        if not current_place:
            continue

        for item in _extract_problem_items_from_report_paragraph(text):
            if item not in communities[current_place]:
                communities[current_place].append(item)

    return OrderedDict(
        (place, _renumber_problem_items(items))
        for place, items in communities.items()
    )


def _extract_problem_items_from_report_paragraph(text: str) -> list[str]:
    if "存在的问题是：" in text:
        return _split_report_problem_items(text.split("存在的问题是：", 1)[1])
    if "存在的问题：" in text:
        return _split_report_problem_items(text.split("存在的问题：", 1)[1])

    for label in ("小区宣传氛围：", "小区公示牌："):
        if label in text:
            return _split_report_problem_items(text.split(label, 1)[1])

    station_match = re.search(r"([0-9A-Za-z一二三四五六七八九十]+号桶站设置情况)：(.+)$", text)
    if station_match:
        station_title = station_match.group(1)
        return [
            f"{station_title}：{item}"
            for item in _split_report_problem_items(station_match.group(2))
        ]
    return []


def _split_report_problem_items(text: str) -> list[str]:
    text = text.strip()
    if is_no_problem_summary(text):
        return []
    parts = split_problem_segments(text)
    if not parts:
        parts = [text]
    items: list[str] = []
    for part in parts:
        item = clean_report_problem_item(part)
        if item and not is_no_problem_summary(item):
            items.append(item)
    return items


def _renumber_problem_items(items: list[str]) -> str:
    if not items:
        return "无问题"
    normalized_items = _drop_duplicate_station_items(
        [_normalize_summary_problem_item(item) for item in items]
    )
    if not normalized_items:
        return "无问题"
    return "；".join(f"（{index}）{item}" for index, item in enumerate(normalized_items, start=1) if item) + "。"


def _normalize_summary_problem_item(text: str) -> str:
    value = clean_problem_number(text)
    value = value.strip("：:，,。；;、 ")
    value = clean_problem_number(value)
    value = value.rstrip("。；;，,、 ")
    return value.strip()


def _drop_duplicate_station_items(items: list[str]) -> list[str]:
    general_categories = {
        _summary_problem_category(item)
        for item in items
        if not _is_station_summary_item(item)
    }
    result: list[str] = []
    for item in items:
        if not item:
            continue
        if _is_station_summary_item(item) and _summary_problem_category(item) in general_categories:
            continue
        result.append(item)
    return result


def _is_station_summary_item(item: str) -> bool:
    return bool(re.search(r"^[0-9A-Za-z一二三四五六七八九十]+号桶站设置情况：", item))


def _summary_problem_category(item: str) -> str:
    value = re.sub(r"^[0-9A-Za-z一二三四五六七八九十]+号桶站设置情况：", "", item)
    value = re.sub(r"\d+处$", "", value)
    value = value.removesuffix("一处").strip("：:，,。；;、 ")
    if "周边" in value and "不洁" in value:
        return "桶站周边不洁"
    if "满冒" in value:
        return "桶站满冒"
    if "站外摆桶" in value or "桶外摆" in value or "垃圾桶外摆" in value:
        return "站外摆桶"
    if "破损" in value or "脏污" in value:
        return "桶站破损、脏污"
    return value


def body_elements(document):
    body = document._body._element
    return list(body)


def remove_elements_between(document, start_index: int, end_index: int) -> None:
    body = document._body._element
    for element in body_elements(document)[start_index:end_index]:
        body.remove(element)


def paragraph_body_index(document, paragraph) -> int:
    for index, element in enumerate(body_elements(document)):
        if element is paragraph._p:
            return index
    raise ValueError("Paragraph is not in document body.")


def set_run_font(run, font_name: str, font_size_pt: int, bold: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size_pt)
    run.bold = bold


def add_formatted_paragraph_at(
    document,
    index: int,
    runs: list[tuple[str, bool]],
    font_name: str,
    font_size_pt: int,
    style=None,
):
    paragraph = document.add_paragraph()
    if style is not None:
        paragraph.style = style
    for text, bold in runs:
        run = paragraph.add_run(text)
        set_run_font(run, font_name, font_size_pt, bold=bold)

    body = document._body._element
    body.remove(paragraph._p)
    body.insert(index, paragraph._p)
    return paragraph


def add_text_paragraph_at(document, index: int, text: str, font_name: str, font_size_pt: int, style=None):
    return add_formatted_paragraph_at(
        document=document,
        index=index,
        runs=[(text, False)],
        font_name=font_name,
        font_size_pt=font_size_pt,
        style=style,
    )


def add_street_heading_at(document, index: int, text: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style = document.styles["Heading 2"] if "Heading 2" in [item.name for item in document.styles] else document.styles["Normal"]
    paragraph = add_text_paragraph_at(document, index, text, FONT_FANGSONG_GB2312, 16, style)
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), "1")
    for run in paragraph.runs:
        set_run_font(run, FONT_FANGSONG_GB2312, 16)
    return paragraph


def add_place_problem_at(document, index: int, place_name: str, problem_text: str):
    return add_formatted_paragraph_at(
        document=document,
        index=index,
        runs=[(place_name, True), (f"{CN_PROBLEM_PREFIX}{problem_text}", False)],
        font_name=FONT_FANGSONG_GB2312,
        font_size_pt=16,
        style=document.styles["Normal"],
    )


def find_paragraph_index(document, predicate) -> int | None:
    for index, paragraph in enumerate(document.paragraphs):
        if predicate(paragraph.text.strip()):
            return index
    return None


def body_index_for_paragraph_number(document, paragraph_number: int) -> int:
    return paragraph_body_index(document, document.paragraphs[paragraph_number])


def default_summary_output_dir(report_date) -> Path:
    return SUMMARY_ROOT / f"\u6bcf\u65e5\u603b\u7ed3\uff08{report_date.short}\uff09"


def chinese_day_prefix(report_date) -> str:
    return f"{report_date.value.month}\u6708{report_date.value.day}\u65e5"


def write_garbage_summary(
    template_path: Path,
    reports: dict,
    report_date,
    output_dir: Path,
    street_report_paths: list[tuple[str, Path]] | None = None,
    ledger_rows: list | None = None,
) -> Path:
    if street_report_paths is not None and _template_has_jinja_tags(template_path):
        try:
            from garbage_daily_report import build_garbage_daily_context, render_garbage_daily_report
        except ModuleNotFoundError:
            from scripts.garbage_daily_report import build_garbage_daily_context, render_garbage_daily_report

        if ledger_rows is not None:
            context = build_garbage_daily_context(ledger_rows, report_date)
            street_context = _summary_context_from_street_reports(street_report_paths, report_date)
            context["residential_streets"] = street_context["residential_streets"]
        else:
            context = _summary_context_from_street_reports(street_report_paths, report_date)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{chinese_day_prefix(report_date)}{CN_GARBAGE_DAILY}.docx"
        return render_garbage_daily_report(template_path.resolve(), context, output_path)
    if ledger_rows is not None:
        try:
            from garbage_daily_report import build_garbage_daily_context, render_garbage_daily_report
        except ModuleNotFoundError:
            from scripts.garbage_daily_report import build_garbage_daily_context, render_garbage_daily_report

        context = build_garbage_daily_context(ledger_rows, report_date)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{chinese_day_prefix(report_date)}{CN_GARBAGE_DAILY}.docx"
        return render_garbage_daily_report(template_path.resolve(), context, output_path)

    from docx import Document
    try:
        import split_street_daily_reports as split
    except ModuleNotFoundError:
        from scripts import split_street_daily_reports as split

    template_path = split.convert_template_to_docx(template_path.resolve())
    document = Document(template_path)
    split.update_template_date(document, report_date)

    start_para = find_paragraph_index(
        document,
        lambda text: text.startswith("\u4e8c\u3001\u533a\u7ea7\u68c0\u67e5\u5c45\u4f4f\u5c0f\u533a/\u5e73\u623f\u80e1\u540c\u60c5\u51b5"),
    )
    if start_para is None:
        raise RuntimeError("Garbage summary template is missing section: 二、区级检查居住小区/平房胡同情况")

    end_para = find_paragraph_index(
        document,
        lambda text: text.startswith("\u4e09\u3001") and "\u533a\u7ea7\u68c0\u67e5\u5c45\u4f4f" not in text,
    )
    if end_para is None or end_para <= start_para:
        end_body_index = len(body_elements(document)) - 1
    else:
        end_body_index = body_index_for_paragraph_number(document, end_para)

    insert_index = body_index_for_paragraph_number(document, start_para) + 1
    remove_elements_between(document, insert_index, end_body_index)

    inserted = 0
    if street_report_paths:
        street_problem_map = street_report_problem_map(street_report_paths)
        checked_street_count = len(street_report_paths)
    else:
        street_problem_map = OrderedDict()
        for street, categories in reports.items():
            places = categories.get(CN_COMMUNITY_CATEGORY, {})
            place_problem_map = OrderedDict()
            for place in places.values():
                problem_text = merge_problem_texts(place_problem_texts(place))
                if problem_text and not is_no_problem_summary(problem_text):
                    place_problem_map[place.name] = problem_text
            if place_problem_map:
                street_problem_map[street] = place_problem_map
        checked_street_count = len(reports)

    intro = (
        f"{report_date.chinese}\uff0c\u533a\u5783\u573e\u5206\u7c7b\u63a8\u8fdb\u5de5\u4f5c\u6307\u6325\u90e8\u529e\u516c\u5ba4"
        f"\u5bf9\u897f\u57ce\u533a{checked_street_count}\u4e2a\u8857\u9053\u751f\u6d3b\u5783\u573e\u5206\u7c7b\u65e5\u5e38\u8fd0\u884c\u60c5\u51b5"
        f"\u8fdb\u884c\u91cd\u70b9\u62bd\u67e5\uff0c\u68c0\u67e5\u5b58\u5728\u7684\u95ee\u9898\u5982\u4e0b\uff1a"
        f"\uff08\u8be6\u89c1\u300a\u897f\u57ce\u533a\u751f\u6d3b\u5783\u573e\u5206\u7c7b\u65e5\u5e38\u8fd0\u884c\u68c0\u67e5\u62a5\u544a-{report_date.chinese}\u300b\uff0c"
        f"\u8bf7\u76f8\u5173\u8857\u9053\u8ba4\u771f\u5206\u6790\uff0c\u5e76\u4e8e3\u65e5\u5185\u5b8c\u6210\u6574\u6539\uff09\u3002"
    )
    add_text_paragraph_at(document, insert_index + inserted, intro, FONT_FANGSONG_GB2312, 16, document.styles["Normal"])
    inserted += 1
    for street_index, (street, places) in enumerate(street_problem_map.items(), start=1):
        add_street_heading_at(document, insert_index + inserted, f"\uff08{chinese_numeral(street_index)}\uff09{street}")
        inserted += 1
        for place_name, problem_text in places.items():
            add_place_problem_at(document, insert_index + inserted, place_name, problem_text)
            inserted += 1

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{chinese_day_prefix(report_date)}{CN_GARBAGE_DAILY}.docx"
    document.save(output_path)
    return output_path


def write_daily_summary(
    template_path: Path,
    reports: dict,
    report_date,
    output_dir: Path,
    ledger_rows: list | None = None,
    street_report_paths: list[tuple[str, Path]] | None = None,
) -> Path:
    if ledger_rows is not None:
        try:
            from docxtpl import DocxTemplate
        except ModuleNotFoundError:
            raise
        try:
            from garbage_daily_report import build_garbage_daily_context
        except ModuleNotFoundError:
            from scripts.garbage_daily_report import build_garbage_daily_context

        context = build_garbage_daily_context(ledger_rows, report_date)
        context["street_problem_counts"] = _street_problem_counts_from_context(context)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{CN_DAILY_SUMMARY}({report_date.value.day}).docx"
        doc = DocxTemplate(str(template_path.resolve()))
        doc.render(context)
        doc.save(str(output_path))
        return output_path
    if street_report_paths is not None:
        try:
            from docxtpl import DocxTemplate
        except ModuleNotFoundError:
            raise

        context = _summary_context_from_street_reports(street_report_paths, report_date)
        context["street_problem_counts"] = _street_problem_counts_from_context(context)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{CN_DAILY_SUMMARY}({report_date.value.day}).docx"
        doc = DocxTemplate(str(template_path.resolve()))
        doc.render(context)
        doc.save(str(output_path))
        return output_path

    from docx import Document
    try:
        import split_street_daily_reports as split
    except ModuleNotFoundError:
        from scripts import split_street_daily_reports as split

    template_path = split.convert_template_to_docx(template_path.resolve())
    document = Document(template_path)
    split.update_template_date(document, report_date)

    marker_para = find_paragraph_index(
        document,
        lambda text: text.startswith("\u4eca\u65e5\u5e02\u7ea7\u68c0\u67e5\u60c5\u51b5\u66f4\u65b0\uff1a"),
    )
    if marker_para is None:
        raise RuntimeError("Daily summary template is missing section: 今日市级检查情况更新：")

    marker_body_index = body_index_for_paragraph_number(document, marker_para)
    remove_elements_between(document, 0, marker_body_index)

    lines = [f"{report_date.chinese}\uff0c\u533a\u5783\u573e\u5206\u7c7b\u6307\u6325\u90e8\u68c0\u67e5{len(reports)}\u4e2a\u8857\u9053\u529e\u4e8b\u5904\uff0c\u5404\u8857\u9053\u95ee\u9898\u6570\u5206\u522b\u4e3a"]
    street_items = list(reports.items())
    for index, (street, categories) in enumerate(street_items):
        places = categories.get(CN_COMMUNITY_CATEGORY, {})
        problem_count = sum(count_real_problems(place_problem_texts(place)) for place in places.values())
        suffix = "\u3002" if index == len(street_items) - 1 else "\uff0c"
        lines.append(f"{short_street_name(street)}{problem_count}\u4e2a{suffix}")

    for offset, line in enumerate(lines):
        add_text_paragraph_at(document, offset, line, FONT_SONG, 12, document.styles["Normal"])

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{CN_DAILY_SUMMARY}({report_date.value.day}).docx"
    document.save(output_path)
    return output_path


def write_summaries(
    reports: dict,
    report_date,
    garbage_template: Path | None = None,
    daily_template: Path | None = None,
    output_dir: Path | None = None,
    street_report_paths: list[tuple[str, Path]] | None = None,
    ledger_rows: list | None = None,
) -> list[Path]:
    output_dir = output_dir or default_summary_output_dir(report_date)
    written: list[Path] = []
    if garbage_template:
        written.append(
            write_garbage_summary(
                garbage_template,
                reports,
                report_date,
                output_dir,
                street_report_paths=street_report_paths,
                ledger_rows=ledger_rows,
            )
        )
    if daily_template:
        written.append(
            write_daily_summary(
                daily_template,
                reports,
                report_date,
                output_dir,
                ledger_rows=ledger_rows,
                street_report_paths=street_report_paths,
            )
        )
    return written


def generate_summaries_from_existing_reports(
    report_dir: Path,
    garbage_template: Path | None = None,
    daily_template: Path | None = None,
    output_dir: Path | None = None,
) -> list[Path]:
    street_report_paths = collect_street_report_paths(report_dir.resolve())
    report_date = infer_report_date_from_street_reports(street_report_paths)
    reports: dict = {}
    print(f"[run] 基于已有日报生成汇总: {report_dir}")
    print(f"[ok] date: {report_date.short}")
    print(f"[ok] existing reports: {len(street_report_paths)}")
    written = write_summaries(
        reports=reports,
        report_date=report_date,
        garbage_template=garbage_template,
        daily_template=daily_template,
        output_dir=output_dir,
        street_report_paths=street_report_paths,
        ledger_rows=None,
    )
    for path in written:
        print(f"[ok] summary: {path}")
    return written


def _template_has_jinja_tags(template_path: Path) -> bool:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return False
    try:
        document = Document(str(template_path))
    except Exception:
        return False
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    text = "\n".join(texts)
    return "{{" in text or "{%" in text


def _summary_context_from_street_reports(street_report_paths: list[tuple[str, Path]], report_date) -> dict:
    try:
        from garbage_daily_report import (
            extract_all_text_from_docx,
            group_outside_bucket_by_street,
            make_outside_bucket_summary,
            parse_outside_bucket_from_street_report_text,
        )
    except ModuleNotFoundError:
        from scripts.garbage_daily_report import (
            extract_all_text_from_docx,
            group_outside_bucket_by_street,
            make_outside_bucket_summary,
            parse_outside_bucket_from_street_report_text,
        )

    street_problem_map = street_report_problem_map(street_report_paths)
    outside_bucket_issues = []
    for _street, path in street_report_paths:
        outside_bucket_issues.extend(parse_outside_bucket_from_street_report_text(extract_all_text_from_docx(path)))
    outside_bucket_by_street = group_outside_bucket_by_street(outside_bucket_issues)
    residential_streets = []
    for index, (street, _) in enumerate(street_report_paths, start=1):
        places = street_problem_map.get(street, OrderedDict())
        street_outside_bucket_issues = outside_bucket_by_street.get(street, [])
        residential_streets.append(
            {
                "index_cn": chinese_numeral(index),
                "name": street,
                "points": [
                    {
                        "name": place_name,
                        "issues": [{"text": item} for item in split_problem_segments(problem_text) if is_real_problem(item)],
                    }
                    for place_name, problem_text in places.items()
                ],
                "outside_bucket_issues": street_outside_bucket_issues,
                "outside_bucket_summary": make_outside_bucket_summary(street, street_outside_bucket_issues),
            }
        )

    previous_day = report_date.value - timedelta(days=1)
    return {
        "report_date": report_date.value,
        "report_date_text": report_date.chinese,
        "street_count": len(street_report_paths),
        "city_check": {
            "summary": "今日市级检查情况未更新。",
            "note": "",
            "has_attachment": False,
            "attachment_no": None,
            "table_title": "市级重点抽查情况",
        },
        "city_check_rows": [],
        "residential_streets": residential_streets,
        "social_streets": [],
        "catering_streets": [],
        "special_check": {"summary": "", "rows": []},
        "enforcement": {
            "attachment_no": 1,
            "date_text": f"{previous_day.year}年{previous_day.month}月{previous_day.day}日",
            "attachment_title": f"{previous_day.month}.{previous_day.day}西城区《北京市生活垃圾管理条例》专线执法统计表",
            "rows": [],
            "total_checks": 0,
            "case_count": 0,
            "fine_amount": 0,
        },
        "waste_data": {
            "date_text": f"{previous_day.year}年{previous_day.month}月{previous_day.day}日",
            "summary": "生活垃圾清运量数据暂未更新。",
        },
    }


def _street_problem_counts_from_context(context: dict) -> list[dict[str, int | str]]:
    counts: OrderedDict[str, int] = OrderedDict()
    for street in context.get("residential_streets", []):
        counts[street["name"]] = sum(
            _issue_count(issue)
            for point in street.get("points", [])
            for issue in point.get("issues", [])
        )
    return [
        {"short_name": short_street_name(street), "problem_count": count}
        for street, count in counts.items()
    ]


def _issue_count(issue: dict) -> int:
    value = issue.get("count")
    if isinstance(value, int):
        return value
    if isinstance(value, str) and value.isdigit():
        return int(value)
    if isinstance(value, str):
        parsed = _parse_chinese_count(value)
        if parsed is not None:
            return parsed
    return problem_segment_count(str(issue.get("text", "")))


def _count_street_problem_map_places(places: "OrderedDict[str, str]") -> int:
    return sum(count_real_problems([problem_text]) for problem_text in places.values())
